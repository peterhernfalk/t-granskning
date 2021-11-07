import docx
from docx import Document
from lxml.etree import _Document

import DOCX_display_document_contents
from DOCX_display_document_contents import *
import glob
import globals
import os
from utilities import write_output, extract_urls_from_table, verify_url_exists

local_test = False

class infomodel_classes:
    def __init__(self, classtitle, document_level, classtable_number):
        self.classtitle = classtitle
        self.document_level = document_level
        self.classtable_number = classtable_number

def IS_inspect_document_contents():
    __set_document_name()
    #document = Document(document_name)
    global document
    __find_all_document_tables()
    __find_all_document_paragraphs()


def IS_inspect_class_description():
    result = True

    ############################################################
    ### DOESN'T WORK!. Compare to document paragraph instead ###
    ############################################################
    for paragraph in range(0, len(class_paragraphs_description)):
        if class_paragraphs_description[paragraph][0:2] == classes_paragraph_level_1+".":
            paragraph_number = class_paragraphs_description[paragraph][0:4].strip()
            if any(x in paragraph_number for x in class_paragraphs_number) == False:
                result = False
                globals.IS_antal_brister_klassbeskrivning += 1

    if result == True:
        write_output("<b>Resultat:</b> alla infomodellklasser har beskrivning")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser har beskrivning")
    else:
        write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar beskrivning")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar beskrivning")

def IS_inspect_attribute_multiplicity():
    all_results = True
    for index in range(len(infomodel_table_indexes)):
        table_num = infomodel_table_indexes[index]
        table_num = index + infomodel_table_indexes[0]
        result = __infomodel_classes_verify_multiplicity(table_num)
        # 2do
        # Insert code from __infomodel_classes_verify_multiplicity
        # 2do
        if result == False:
            tbl_no = index + infomodel_table_indexes[0]
            table_title = __get_infomodel_classname_from_table_number(tbl_no, True)
            write_output(globals.HTML_3_SPACES + "Infomodellklass " + table_title + ": saknar multiplicitet")
            write_detail_box_content(globals.HTML_3_SPACES + "Infomodellklass " + table_title + ": saknar multiplicitet")
            globals.IS_antal_brister_multiplicitet += 1
            all_results = False
    if all_results == True:
        write_output("<b>Resultat:</b> alla infomodellklasser har multiplicitet i datatypskolumnen")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser har multiplicitet i datatypskolumnen")
    else:
        write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar multiplicitet i datatypskolumnen")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar multiplicitet i datatypskolumnen")

### Data Types ###
def IS_inspect_usage_of_defined_datatypes():
    all_results = True
    table = document.tables[datatype_definitions_table[0]]
    defined_datatypes = []
    for i in range(1,len(table.rows)):
        defined_datatypes.append(table.cell(i, 0).text.strip().lower())

    for index in range(len(infomodel_table_indexes)):
        table_number = infomodel_table_indexes[index]
        table = document.tables[table_number]
        table_column_index = __get_table_column_index(table_number, "datatyp")
        if table_column_index == 0:
            table_column_index = __get_table_column_index(table_number, "data-typ")

        for i in range(1, len(table.rows)):
            if any(x in table.cell(i, table_column_index).text.strip().lower() for x in defined_datatypes) == False:
                table_title = __get_title_by_table_number(index)
                write_output(globals.HTML_3_SPACES + "Datatypen är ej definierad!  Tabell: " +
                        table_title + ". Attribut: " + table.cell(i, 0).text + ". Datatyp: " +
                        table.cell(i, table_column_index).text.strip().lower())
                write_detail_box_content(globals.HTML_3_SPACES + "Datatypen är ej definierad!  Tabell: " +
                        table_title + ". Attribut: " + table.cell(i, 0).text + ". Datatyp: " +
                        table.cell(i, table_column_index).text.strip().lower())
                globals.IS_antal_brister_datatyper += 1
                all_results = False
    if all_results == True:
        write_output("<b>Resultat:</b> alla infomodellklasser använder definierade datatyper")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser använder definierade datatyper")
    else:
        write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar eller använder icke definierad datatyp")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar eller använder icke definierad datatyp")

def IS_inspect_usage_of_reference_infomodel():
    all_results = True
    table = document.tables[datatype_definitions_table[0]]

    for index in range(len(infomodel_table_indexes)):
        table_num = infomodel_table_indexes[index]
        table = document.tables[table_num]
        for i in range(1, len(table.rows)):
            if table.cell(i, 1).text.strip() == "":
                tbl_no = index+infomodel_table_indexes[0]
                table_title = __get_infomodel_classname_from_table_number(tbl_no, True)
                write_output(globals.HTML_3_SPACES + "Mappning saknas till RIM!  Tabell: " + table_title + ". Attribut: " + table.cell(i, 0).text)
                write_detail_box_content(globals.HTML_3_SPACES + "Mappning saknas till RIM!  Tabell: " + table_title + ". Attribut: " + table.cell(i, 0).text)
                globals.IS_antal_brister_referensinfomodell += 1
                all_results = False
    if all_results == True:
        write_output("<b>Resultat:</b> alla infomodellklasser är mappade till referensinformationsmodellen")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser är mappade till referensinformationsmodellen")
    else:
        write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar mappning till referensinformationsmodellen")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar mappning till referensinformationsmodellen")


################################################## local methods
def __get_infomodel_classname_from_table_number(table_number, include_level):
    global infomodel_classes_list
    result_classtitle = ""
    for obj in infomodel_classes_list:
        if obj.classtable_number == table_number:
            if include_level == True:
                result_classtitle = obj.document_level + " " + obj.classtitle
            else:
                result_classtitle = obj.classtitle
            break
    return result_classtitle

def __get_infomodel_table_number_from_classname(classtitle):
    global infomodel_classes_list
    result_table_number = ""
    for obj in infomodel_classes_list:
        classtitle_without_level = obj.classtitle.strip().lower()
        if classtitle.strip().lower() == classtitle_without_level:
            result_table_number = obj.classtable_number
            break
    return result_table_number

def __set_document_name():
    global document_name
    """os.chdir(globals.document_path)
    for word_document in glob.glob("IS_*.doc*"):
        document_name = r""+globals.document_path+"/"+word_document"""
    document_name = globals.IS_document_name

### Find all tables with infomodel class information ###
def __find_all_document_tables():
    global all_tables
    global datatype_definitions_table
    global document
    global infomodel_table_indexes
    #document = Document(document_name)
    all_tables = document.tables
    infomodel_table_indexes = []
    datatype_definitions_table = []
    index = 0
    for table in all_tables:
        if table.cell(0,0).text.strip().lower() == "attribut":
            infomodel_table_indexes.append(index)
        if table.cell(0, 0).text.strip().lower() == "förkortning":
            datatype_definitions_table.append(index)
        elif table.cell(0,0).text.strip().lower() == "benämning":
            datatype_definitions_table.append(index)
        index += 1

### Find and save a list of the paragraphs ###
def __find_all_document_paragraphs():
    level_from_style_name = {f'Rubrik {i} Nr': i for i in range(10)}
    current_levels = [0] * 10

    global all_paragraphs
    global classes_paragraph_level_1
    global class_paragraphs_number
    global class_paragraphs_title
    global class_paragraphs_description
    all_paragraphs = []
    class_paragraphs_number = []
    class_paragraphs_title = []
    class_paragraphs_description = []
    classes_paragraph_level_1 = ""

    for paragraph in document.paragraphs:
        if paragraph.style.name not in level_from_style_name:
            all_paragraphs.append(paragraph.text)
        else:
            level = level_from_style_name[paragraph.style.name]
            current_levels[level] += 1
            #
            class_paragraphs_description.append(__format_levels(current_levels) + ' ' + paragraph.text)
            #
            for level in range(level + 1, 10):
                current_levels[level] = 0
            if paragraph.text[0:20].strip().lower() == "klasser och attribut":
                classes_paragraph_level_1 = __format_levels(current_levels)

            # Bug: missing paragraph descriotion => doesn't append the paragraph
            if __format_levels(current_levels)[0] == classes_paragraph_level_1:
                if __format_levels(current_levels)[0:2] == classes_paragraph_level_1 + ".":
                    class_paragraphs_title.append(__format_levels(current_levels) + ' ' + paragraph.text)
                    class_paragraphs_number.append(__format_levels(current_levels))

            all_paragraphs.append(__format_levels(current_levels) + ' ' + paragraph.text)

def __get_table_column_index(table_number, searchphrase):
    table = document.tables[table_number]
    column_count = len(table.row_cells(0))
    column_index = 0
    for i in range(0,column_count):
        if searchphrase.strip().lower() == "multiplicitet":
            if "multiplicitet" in table.cell(0, i).text.lower() or "multi-plicitet" in table.cell(0, i).text.lower():
                column_index = i
                break
        else:
            if searchphrase.strip().lower() in table.cell(0, i).text.lower():
                column_index = i
                break
    return column_index

def __get_title_by_table_number(table_number):
    table_title = ""

    ###############################################
    ### 2do: get paragraph title for this table ###
    tbl_no = table_number + infomodel_table_indexes[0]
    table_title = __get_infomodel_classname_from_table_number(tbl_no, True)
    if table_title.strip() == "":
        table_title = str(table_number)
    ###############################################

    return table_title

def __infomodel_classes_verify_multiplicity(table_number):
    result = True
    table = document.tables[table_number]
    allowed_multiplicities = ['0', '1', '*']
    table_column_index = __get_table_column_index(table_number, "multiplicitet")
    if table_column_index == 0:
        table_column_index = __get_table_column_index(table_number, "datatyp")
    if table_column_index == 0:
        table_column_index = __get_table_column_index(table_number, "data-typ")

    for i in range(1,len(table.rows)):
        if any(cell_multiplicity in table.cell(i, table_column_index).text for cell_multiplicity in allowed_multiplicities) == False:
            tbl_no = table_number #+ infomodel_table_indexes[0]
            table_title = __get_infomodel_classname_from_table_number(tbl_no, True)
            write_output(globals.HTML_3_SPACES + globals.HTML_3_SPACES + "Klass: " + table_title + " saknar multiplicitet för: " + table.cell(i, 0).text)
            result = False
    return result

def __inspect_revision_history():
    table = document.tables[1]
    # 2do: Display version number from page 1 in the document

    if table.cell(0,0).text == "Revisionshistorik mall":
        table = document.tables[2]

    for i, row in enumerate(table.rows):
        text = tuple(cell.text for cell in row.cells)

    if str(table.cell(i, 0).text) != globals.tag:
        write_output("OBS! Revisionshistoriken behöver uppdateras. (hittade: "+str(table.cell(i, 0).text)+" men förväntade: "+globals.tag+")")
    else:
        write_output("Revisionshistoriken är uppdaterad för denna version av domänen")
    write_output("Revisionshistorikens sista rad: " + str(text))

def __format_levels(current_level):
    levels = [str(level) for level in current_level if level != 0]
    return '.'.join(levels)  # Customize your format here

def __inspect_classname_format():
    result = True

    ### 2do: check that the class name begins with UPPERCASE ###
    write_output("<br>Kontroll att infomodellklassernas rubriker är i alfabetisk ordning och börjar med stor bokstav")
    for paragraph in range(0, len(class_paragraphs_title)):
        write_output(globals.HTML_3_SPACES + class_paragraphs_title[paragraph])
    write_output("För närvarande sker kontrollen manuellt, med ovanstående listning som underlag")

    return result

def __inspect_class_order():
    ### 2do ###
    result = True

def __infomodel_classes_verify_attribute_case(table_number):
    table = document.tables[table_number]
    result = True
    for i in range(1,len(table.rows)):
        if table.cell(i, 0).text[0:1].islower() == False:
            write_output(globals.HTML_3_SPACES + "Fel skiftläge för första bokstaven i: " + table.cell(i, 0).text)
            result = False
    return result

def __inspect_attribut_name_case():
    all_results = True
    write_output("\nKontroll att attributnamn inleds med liten bokstav i infomodellklassernas tabeller")
    for index in range(len(infomodel_table_indexes)):
        table_num = infomodel_table_indexes[index]
        result = __infomodel_classes_verify_attribute_case(table_num)
        # 2do
        # Insert code from __infomodel_classes_verify_attribute_case
        # 2do
        if result == False:
            write_output(globals.HTML_3_SPACES + "Infomodellklass " + str(index+1) + ": innehåller attributnamn med fel skiftläge för första bokstaven\n")
            all_results = False
    if all_results == True:
        write_output("<b>Resultat:</b> alla infomodellklasser är OK avseende skiftläge på attributnamnens första bokstav")
    else:
        write_output("<b>Resultat:</b> en eller flera infomodellklasser har fel skiftläge för första bokstaven i attribut")


############################## TEST ##############################
def __extract_hyperlink(xml):
    import re
    # text = xml.decode('utf-8')
    text = xml
    if "</w:hyperlink>" in text:
        print(text)
    #text = text.replace("</w:hyperlink>", "")
    #text = re.sub('<w:hyperlink[^>]*>', "", text)
    # return text.encode('utf-8')
    #return text

if local_test == True:
    globals.document_path = "/Users/peterhernfalk/Desktop/Aktuellt/_T-granskningar/git-Repo/riv.clinicalprocess.healthcond.certificate/docs"
    #globals.document_path = "/Users/peterhernfalk/Desktop/Aktuellt/_T-granskningar/git-Repo/riv.informationsecurity.authorization.consent/docs"

    #DOCX_prepare_inspection("IS_*.doc*")
    __set_document_name()
    document = Document(document_name)
    __find_all_document_tables()
    __find_all_document_paragraphs()
    #IS_inspect_attribute_multiplicity()
    #IS_inspect_usage_of_defined_datatypes()
    #DOCX_inspect_revision_history()
    #DOCX_display_paragraph_text_and_tables("klasser och attribut", True, False)
    #DOCX_display_paragraph_text_and_tables("processmodell", True, False)

    #IS_inspect_reference_links()

def IS_find_empty_table_cells():
    result = False
    for table_index in range(len(infomodel_table_indexes)):
        table_number = infomodel_table_indexes[table_index]
        table = document.tables[table_number]
        for row in range(1, len(table.rows)):
            column_count = len(table.row_cells(0))
            for column in range(0, column_count):
                if table.cell(row, column).text.strip() == "":
                    result = True
                    #table_title = __get_title_by_table_number(table_index + 1)
                    tbl_no = table_index + infomodel_table_indexes[0]
                    table_title = __get_infomodel_classname_from_table_number(tbl_no, True)
                    write_output(globals.HTML_3_SPACES + "Tabellcell utan innehåll funnen!  Tabell: " + str(table_title) + ", Rad: " + str(row) + ", Kolumn: " + str(column+1))
                    write_detail_box_content(globals.HTML_3_SPACES + "Tabellcell utan innehåll funnen!  Tabell: " + str(table_title) + ", Rad: " + str(row) + ", Kolumn: " + str(column+1))
                    globals.IS_antal_brister_tomma_tabellceller += 1

    if result == True:
        write_output("<b>Resultat:</b> det finns infomodellklass(er) med en eller flera celler utan innehåll")
        write_detail_box_content("<b>Resultat:</b> det finns infomodellklass(er) med en eller flera celler utan innehåll")
    else:
        write_output("<b>Resultat:</b> alla infomodellklassers alla celler har innehåll")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklassers alla celler har innehåll")

    if local_test == True:
        print("\ndoamin:", globals.domain)
        #author = document.core_properties.author
        #print("document.core_properties.author:",author)
        print("class_paragraphs_title:", class_paragraphs_title)
        print("class_paragraphs_number:", class_paragraphs_number)
        print("infomodel_table_indexes:",infomodel_table_indexes)

        print("\ninfomodel_classes_list:")
        #global infomodel_classes_list
        #infomodel_classes_list = []
        #IS_init_infomodel_classes_list()

        for obj in infomodel_classes_list:
            print("\t", obj.document_level, obj.classtitle, obj.classtable_number, sep=' ')
        #print("Infomodellklass för tabell 12:", __get_infomodel_classname_from_table_number(12))
        #print("Infomodellklass för tabell 29:", __get_infomodel_classname_from_table_number(29))
        #print("Infomodelltabell för delsvar:", __get_infomodel_table_number_from_classname("delsvar"))
        #print("Infomodelltabell för svar:", __get_infomodel_table_number_from_classname("svar"))
        #print("Infomodelltabell saknas för avsändare:", __get_infomodel_table_number_from_classname("avsändare"))
        #print("Infomodelltabell saknas för mottagare:", __get_infomodel_table_number_from_classname("mottagare"))


#################################################################################################
#################################################################################################
def IS_init_infomodel_classes_list():
    global document
    global document_name
    __set_document_name()
    #print("document_name:",document_name)
    #print("globals.docx_IS_document",globals.docx_IS_document)
    #document = Document(document_name)
    #document = Document(globals.IS_document_name)
    document = globals.docx_IS_document
    #print("document",document)
    __find_all_document_paragraphs()
    __find_all_document_tables()

    global infomodel_classes_list
    infomodel_classes_list = []
    searched_paragraph_level = 2    #2do: replace with correct value
    append_paragraph_title = ""
    paragraph_level = ""
    table_number = infomodel_table_indexes[0]

    for block in __iter_block_items(document, searched_paragraph_level):
        append_table_number = 0
        if isinstance(block, Paragraph):
            new_paragraph_level = DOCX_document_structure_get_levelvalue(block.text)
            if new_paragraph_level != NOT_FOUND:
                paragraph_level = new_paragraph_level
            if len(paragraph_level) == 1:
                continue
            if paragraph_level != "" and paragraph_level[0] == classes_paragraph_level_1[0]:
                if DOCX_document_structure_get_levelvalue(block.text) != NOT_FOUND:
                    append_paragraph_title = block.text
        elif isinstance(block, Table):
            if paragraph_level != "":
                if paragraph_level[0] == classes_paragraph_level_1[0]:
                    append_table_number = table_number
                    table_number += 1

        if append_paragraph_title != "" and append_table_number != 0:
            #print("\tappend:", append_paragraph_title, paragraph_level, append_table_number)
            append_paragraph_title = append_paragraph_title.replace("\n"," ")
            infomodel_classes_list.append(infomodel_classes(append_paragraph_title, paragraph_level, append_table_number))
            append_paragraph_title == ""

        #for obj in infomodel_classes_list:
        #    print("\t", obj.document_level, obj.classtitle, obj.classtable_number, sep=' ')
        #print("Infomodellklass för tabell 12:", __get_infomodel_classname_from_table_number(12,False))

    #pet: test
    IS_init_all_tables_list()
    #pet: test


def IS_init_all_tables_list():
    global document
    #__set_document_name()
    #document = Document(document_name)
    #__find_all_document_paragraphs()
    #__find_all_document_tables()

    global all_tables_list
    all_tables_list = []
    searched_paragraph_level = 1    #2do: replace with correct value
    append_paragraph_title = ""
    paragraph_level = ""
    table_number = 0    #infomodel_table_indexes[0]

    for block in __iter_block_items(document, searched_paragraph_level):
        append_table_number = 0
        if isinstance(block, Paragraph):
            new_paragraph_level = DOCX_document_structure_get_levelvalue(block.text)
            if new_paragraph_level != NOT_FOUND:
                paragraph_level = new_paragraph_level
            if len(paragraph_level) == 1:
                continue
            if paragraph_level != "": #and paragraph_level[0] == classes_paragraph_level_1[0]:
                if DOCX_document_structure_get_levelvalue(block.text) != NOT_FOUND:
                    append_paragraph_title = block.text
        elif isinstance(block, Table):
            if paragraph_level != "":
                #if paragraph_level[0] == classes_paragraph_level_1[0]:
                #    append_table_number = table_number
                #    table_number += 1
                append_table_number = table_number
                table_number += 1

        if append_paragraph_title != "" and append_table_number != 0:
            #print("\tappend:", append_paragraph_title, paragraph_level, append_table_number)
            append_paragraph_title = append_paragraph_title.replace("\n"," ")
            all_tables_list.append(infomodel_classes(append_paragraph_title, paragraph_level, append_table_number))
            append_paragraph_title == ""
    #print("for obj in all_tables_list")
    #for obj in all_tables_list:
    #    print("\t", obj.document_level, obj.classtitle, obj.classtable_number, sep=' ')
    #print("Infomodellklass för tabell 12:", __get_infomodel_classname_from_table_number(12,False))


def __document_structure_get_key(searched_value):
    for key, value in DOCX_display_document_contents.document_structure_dict.items():
        if searched_value == value:
            return key
    return NOT_FOUND

def __iter_block_items(parent,searched_paragraph_level):
    if isinstance(parent, docx.document.Document):   #_Document
        parent_elm = parent.element.body

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
#################################################################################################
#################################################################################################
