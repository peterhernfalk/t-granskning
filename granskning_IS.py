
import datetime
import document_mangagement
import docx
from docx.table import *
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import *
from docx.oxml.table import *
from docx.api import Document  # noqa
import DOCX_display_document_contents
from utilities import *


IS_antal_brister_attributnamn = 0
IS_antal_brister_datatyper = 0
IS_antal_brister_klassbeskrivning = 0
IS_antal_brister_multiplicitet = 0
IS_antal_brister_referensinfomodell = 0
IS_antal_brister_referenslänkar = 0
IS_antal_brister_revisionshistorik = 0
IS_antal_brister_tomma_begreppsbeskrivningstabellceller = 0
IS_antal_brister_tomma_referenstabellceller = 0
IS_antal_brister_tomma_revisionshistoriktabellceller = 0
IS_antal_brister_tomma_tabellceller = 0
IS_begreppslista_finns = False
IS_begreppsmodell_finns = False
IS_detail_box_contents = ""
IS_document_exists = False
IS_document_name = ""
IS_exists = False
IS_felmeddelande = ""
IS_informationsmodell_finns = False
IS_kodverkstabell_finns = False
IS_referensinfomodell_finns = False

def prepare_IS_inspection(domain, tag, alt_document_name):
    print("IS-init påbörjas",datetime.datetime.now().replace(microsecond=0))
    """
    Beräknar url till infospecdokumentet för angiven domain och tag.

    Laddar ner dokumentet till en virtuell fil som läses in i ett docx-Document.

    Anropar därefter metoden "INFO_inspect_document" som genomför granskning av dokumentet.
    """
    """
    2do: Förenkla och snygga till koden
    """

    global IS_antal_brister_attributnamn
    global IS_antal_brister_datatyper
    global IS_antal_brister_klassbeskrivning
    global IS_antal_brister_multiplicitet
    global IS_antal_brister_referensinfomodell
    global IS_antal_brister_referenslänkar
    global IS_antal_brister_revisionshistorik
    global IS_antal_brister_tomma_begreppsbeskrivningstabellceller
    global IS_antal_brister_tomma_referenstabellceller
    global IS_antal_brister_tomma_revisionshistoriktabellceller
    global IS_antal_brister_tomma_tabellceller
    global IS_begreppslista_finns
    global IS_begreppsmodell_finns
    global IS_detail_box_contents
    global IS_document_exists
    global IS_document_name
    global IS_exists
    global IS_felmeddelande
    global IS_informationsmodell_finns
    global IS_kodverkstabell_finns
    global IS_referensinfomodell_finns
    IS_antal_brister_attributnamn = 0
    IS_antal_brister_datatyper = 0
    IS_antal_brister_klassbeskrivning = 0
    IS_antal_brister_multiplicitet = 0
    IS_antal_brister_referensinfomodell = 0
    IS_antal_brister_referenslänkar = 0
    IS_antal_brister_revisionshistorik = 0
    IS_antal_brister_tomma_begreppsbeskrivningstabellceller = 0
    IS_antal_brister_tomma_referenstabellceller = 0
    IS_antal_brister_tomma_revisionshistoriktabellceller = 0
    IS_antal_brister_tomma_tabellceller = 0
    IS_begreppslista_finns = False
    IS_begreppsmodell_finns = False
    IS_detail_box_contents = ""
    IS_document_exists = False
    IS_document_name = ""
    IS_exists = False
    IS_informationsmodell_finns = False
    IS_kodverkstabell_finns = False
    IS_referensinfomodell_finns = False
    IS_felmeddelande = ""

    global IS_page_link
    IS_page_link = document_mangagement.DOC_get_document_page_link(domain, tag, globals.IS)
    downloaded_IS_page = document_mangagement.DOC_get_downloaded_document(IS_page_link)

    IS_head_hash = document_mangagement.DOC_get_head_hash(downloaded_IS_page)
    IS_document_link = document_mangagement.DOC_get_document_link(domain, tag, globals.IS, IS_head_hash, alt_document_name)
    downloaded_IS_document = document_mangagement.DOC_get_downloaded_document(IS_document_link)
    if downloaded_IS_document.status_code == 404:
        IS_exists = False
    else:
        globals.docx_IS_document = document_mangagement.DOC_get_docx_document(downloaded_IS_document)
        IS_document_exists = True
        IS_exists = True

        DOCX_display_document_contents.DOCX_prepare_inspection("IS_*.doc*")
        IS_init_infomodel_classes_list()


def perform_IS_inspection(domain, tag, alt_document_name):
    prepare_IS_inspection(domain, tag, alt_document_name)
    print("IS-granskning påbörjas",datetime.datetime.now().replace(microsecond=0))

    global IS_antal_brister_attributnamn
    global IS_antal_brister_datatyper
    global IS_antal_brister_klassbeskrivning
    global IS_antal_brister_multiplicitet
    global IS_antal_brister_referensinfomodell
    global IS_antal_brister_referenslänkar
    global IS_antal_brister_revisionshistorik
    global IS_antal_brister_tomma_begreppsbeskrivningstabellceller
    global IS_antal_brister_tomma_referenstabellceller
    global IS_antal_brister_tomma_revisionshistoriktabellceller
    global IS_antal_brister_tomma_tabellceller
    global IS_begreppslista_finns
    global IS_begreppsmodell_finns
    global IS_detail_box_contents
    global IS_document_exists
    global IS_document_name
    global IS_exists
    global IS_felmeddelande
    global IS_informationsmodell_finns
    global IS_kodverkstabell_finns
    global IS_referensinfomodell_finns

    if IS_exists == False:
        return

    write_detail_box_content("<b>Krav:</b> om dokumentegenskaper finns ska version och ändringsdatum stämma överens med granskad version")
    # 2do: kontrollera dokumentegenskaper avseende versionsnummer   https://python-docx.readthedocs.io/en/latest/dev/analysis/features/coreprops.html
    """
        Exempel på Core properties:
            Title, Subject, Author

        Exempel på Custom properties:
            datepublished, datumpubliserad, Publisheddate
            domain_1,_2,_3
            svekortnamn
            svename, svenamn, SvensktDomänNamn
            Version, version, Version_1,_2,_3, Version_RC, version1,2,3
    """
    # 2do: kontrollera versionsnummer på dokumentets första sida: förekomst av "Version" med efterföljande versionsnummer

    print("\tIS: revisionshistorik, version",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> revisionshistoriken ska vara uppdaterad för samma version som domänen")
    write_detail_box_content("<b>Granskningsstöd:</b> om revisionshistoriken inte är uppdaterad, kontakta beställaren eller skriv en granskningskommentar")

    used_table_no = DOCX_display_document_contents.DOCX_get_tableno_for_paragraph_title("revisionshistorik")
    if used_table_no > 0:
        IS_antal_brister_revisionshistorik = DOCX_display_document_contents.DOCX_inspect_revision_history(globals.IS,used_table_no)
    else:
        IS_antal_brister_revisionshistorik = DOCX_display_document_contents.DOCX_inspect_revision_history(globals.IS,globals.TABLE_NUM_REVISION)

    print("\tIS: revisionshistorik, tomma celler",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> revisionshistorikens alla tabellceller ska ha innehåll")
    if used_table_no > 0:
        result, IS_antal_brister_tomma_revisionshistoriktabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(used_table_no, True, globals.DISPLAY_TYPE_TABLE)
    else:
        result, IS_antal_brister_tomma_revisionshistoriktabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(globals.TABLE_NUM_REVISION, True, globals.DISPLAY_TYPE_TABLE)

    print("\tIS: referenstabellslänkar",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> länkarna i referenstabellen ska fungera")
    used_table_no = DOCX_display_document_contents.DOCX_get_tableno_for_paragraph_title("referenser")
    links_excist = False
    if used_table_no > 0:
        links_excist, IS_antal_brister_referenslänkar = DOCX_display_document_contents.DOCX_inspect_reference_links(used_table_no)
    else:
        links_excist, IS_antal_brister_referenslänkar = DOCX_display_document_contents.DOCX_inspect_reference_links(globals.TABLE_NUM_REF)
    if IS_antal_brister_referenslänkar > 0:
        write_detail_box_content("<b>Resultat:</b> en eller flera länkar är felaktiga, eller kan inte tolkas korrekt av granskningsfunktionen.")
        write_detail_box_content("<b>Granskningsstöd:</b> gör manuell kontroll i dokumentet av de länkar som rapporteras som felaktiga")
    else:
        if links_excist == True:
            write_detail_box_content("<b>Resultat:</b> alla kontrollerade länkar fungerar")
        else:
            write_detail_box_content("<b>Resultat:</b> inga länkar har kontrollerats")

    print("\tIS: referenstabell, tomma celler",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> referenstabellens alla tabellceller ska ha innehåll")
    if used_table_no > 0:
        result, IS_antal_brister_tomma_referenstabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(used_table_no, True, globals.DISPLAY_TYPE_TEXT)
    else:
        result, IS_antal_brister_tomma_referenstabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(globals.TABLE_NUM_REF, True, globals.DISPLAY_TYPE_TEXT)

    print("\tIS: referensmodell, förekomst och innehåll",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> Referensmodellsförteckning ska finnas och ha innehåll")
    write_detail_box_content("<b>Krav:</b> Versionskolumnen ska finnas och ha innehåll")
    # 2do: kontrollera att det finns innehåll i referensmodelltabellens versionskolumn, Kolumnrubrik: "Version"
    IS_referensinfomodell_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables("Referensmodellsförteckning (RIM)", globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.TABLES)
    if IS_referensinfomodell_finns == False:
        write_detail_box_content("<b>Granskningsstöd:</b> inget innehåll visas, vilket kan bero på att avsnittsrubriken saknas eller är annan än den förväntade (Referensmodellsförteckning (RIM))")
    write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående listning som underlag")

    print("\tIS: begreppsmodell ska finnas",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infospecen ska innehålla ett avsnitt för begreppsmodell och beskrivning av begrepp")
    IS_begreppsmodell_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables("Begreppsmodell och beskrivning", globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.NO_TABLES)
    write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående listning som underlag")

    print("\tIS: begreppsbeskrivningstabell ska ha innehåll",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> begreppsmodellens tabell med begreppsbeskrivningar ska finnas och ha innehåll")
    used_table_no = DOCX_display_document_contents.DOCX_get_tableno_for_paragraph_title("begreppsmodell och beskrivning")
    if used_table_no > 0:
        begreppsbeskrivning_tabell = used_table_no
    else:
        begreppsbeskrivning_tabell = DOCX_display_document_contents.DOCX_get_tableno_for_first_column_title("begrepp", document.tables)
    result, IS_antal_brister_tomma_begreppsbeskrivningstabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(begreppsbeskrivning_tabell, True, globals.DISPLAY_TYPE_TABLE)

    print("\tIS: begreppslista ska finnas",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infospecen ska innehålla en begreppslista")
    IS_begreppslista_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables("Begreppssystem, klassifikationer och kodverk", globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.NO_TABLES)
    if IS_begreppslista_finns == False:
        write_detail_box_content("<b>Granskningsstöd:</b> inget innehåll visas, vilket kan bero på att avsnittsrubriken saknas eller är annan än den förväntade (Begreppssystem, klassifikationer och kodverk)")
    write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående avsnittsinnehåll som underlag")

    # 2do: kontrollera att begrepp i begreppbeskrivningstabellen finns definierade i dokumentets begreppslista
    # tolkning: jämför begreppskolumnen med beskrivningskolumnen i begreppsbeskrivningstabellen
    print("\tIS: begrepp ska finas definierade i begreppslistan",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> begrepp i begreppsbeskrivningstabellen ska finnas definierade i dokumentets begreppslista")
    if begreppsbeskrivning_tabell > 0 and IS_begreppslista_finns == True:
        write_detail_box_content("<b>Resultat:</b> kontrollen är inte utvecklad än, så för närvarande kan inget resultat visas!")
    else:
        write_detail_box_content("<b>Resultat:</b> kravet är inte uppfyllt eftersom inte både begreppsbeskrivningstabellen och begreppslistan finns i dokumentet")

    print("\tIS: infomodell ska finnas",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infospecen ska innehålla ett avsnitt för Informationsmodell")
    IS_informationsmodell_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables("Informationsmodell och beskrivning", globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.NO_TABLES)
    if IS_informationsmodell_finns == False:
        write_detail_box_content("<b>Granskningsstöd:</b> inget innehåll visas, vilket kan bero på att avsnittsrubriken saknas eller är annan än den förväntade (Informationsmodell och beskrivning)")
    write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående listning som underlag")

    print("\tIS: lista infomodellklasser",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br><b>Krav:</b> infomodellklasserna ska komma i alfabetisk ordning")
    write_detail_box_content("<b>Krav:</b> infomodellklassernas rubriker ska börja med stor bokstav")
    write_detail_box_content("<b>Granskningsstöd:</b> kontrollera att infomodellklassernas rubriker är i alfabetisk ordning")
    DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables("klasser och attribut", globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.NO_TABLES)
    paragraph_title_list, antal_klasser_liten_begynnelsebokstav = DOCX_display_document_contents.DOCX_list_searched_paragraph_titles_wrong_case("klasser och attribut", "Klass ", globals.UPPER_CASE)
    # print("antal_klasser_liten_begynnelsebokstav",antal_klasser_liten_begynnelsebokstav,"\nparagraph_title_list",paragraph_title_list)
    write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående listning som underlag")
    ##IS_inspect_document_contents()

    print("\tIS: infomodellklasser, attributnamn",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> Infomodellklassernas attributnamn ska ha liten begynnelsebokstav")
    IS_antal_brister_attributnamn = IS_inspect_attribute_case()

    print("\tIS: infomodellklasser, beskrivning",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infomodellklassernas rubriker ska ha beskrivning i anslutning till rubriken")
    IS_antal_brister_klassbeskrivning = IS_inspect_class_description()

    print("\tIS: infomodellklasser, multiplicitet",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> multiplicitet ska vara ifyllt i infomodellklassernas tabeller")
    IS_antal_brister_multiplicitet = IS_inspect_attribute_multiplicity()

    print("\tIS: infomodellklasser, definierade datatyper",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infomodellklassernas attribut ska använda definierade datatyper")
    IS_antal_brister_datatyper = IS_inspect_usage_of_defined_datatypes()

    ### 2do ### Anropa DOCX_get_tableno_for_paragraph_title för att få reda på om kodverkstabell finns
    print("\tIS: kodverkstabell ska finnas",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infospecen ska innehålla en tabell med användna kodverk")
    search_phrase_kodverk = "Identifikationer och kodverk"
    IS_kodverkstabell_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables(search_phrase_kodverk, globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.NO_TABLES)
    if IS_kodverkstabell_finns == False:
        search_phrase_kodverk = "Identifierare och kodverk"
        IS_kodverkstabell_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables(search_phrase_kodverk, globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.NO_TABLES)
        if IS_kodverkstabell_finns == False:
            search_phrase_kodverk = "Begreppssystem, klassifikationer och kodverk"
            IS_kodverkstabell_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables(search_phrase_kodverk, globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.NO_TABLES)
            if IS_kodverkstabell_finns == False:
                write_detail_box_content("<b>Granskningsstöd:</b> inget av avsnitten 'Identifikationer och kodverk' eller 'Begreppssystem, klassifikationer och kodverk' hittades i infospecen")
    write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående listning som underlag")
    # 2do: jämför klasstabellernas med dokumentets kodverkstabell
    # Kolumner: data, kodverk/format/regler, kodverk

    print("\tIS: kodverkstabellens innehåll",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> Kodverkstabellen ska ha relevant innehåll")
    if IS_kodverkstabell_finns == True:
        DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables(search_phrase_kodverk, globals.TITLE, globals.NO_INITIAL_NEWLINE, globals.NO_TEXT, globals.TABLES)
    write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående listning som underlag")

    print("\tIS: infomodellklasser, mappning till RIM",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infomodellklassernas attribut ska vara mappade till referensinformationsmodellen")
    IS_antal_brister_referensinfomodell = IS_inspect_usage_of_reference_infomodel()

    print("\tIS: infomodellklasser, tomma celler",datetime.datetime.now().replace(microsecond=0))
    write_detail_box_content("<br>")
    write_detail_box_content("<b>Krav:</b> infomodellklassernas alla celler ska innehålla värde")
    empty_cells_found = False
    antal_tomma_klasstabellceller = 0
    for table_index in range(len(infomodel_table_indexes)):
        table_number = infomodel_table_indexes[table_index]
        result, antal = DOCX_display_document_contents.DOCX_empty_table_cells_exists(table_number, False, globals.DISPLAY_TYPE_TEXT)
        if result == True:  # DOCX_empty_table_cells_exists(table_number, False, globals.DISPLAY_TYPE_TEXT)
            empty_cells_found = True
            # antal_tomma_klasstabellceller += globals.IS_antal_brister_tomma_tabellceller
            antal_tomma_klasstabellceller += antal
    if empty_cells_found == True:
        write_detail_box_content("<b>Resultat:</b> det finns infomodellklass(er) med en eller flera celler utan innehåll")
    else:
        write_detail_box_content("<b>Resultat:</b> alla infomodellklassers alla celler har innehåll")
    IS_antal_brister_tomma_tabellceller = antal_tomma_klasstabellceller
    print("IS-granskning klar",datetime.datetime.now().replace(microsecond=0))



class infomodel_classes:
    def __init__(self, classtitle, document_level, classtable_number):
        self.classtitle = classtitle
        self.document_level = document_level
        self.classtable_number = classtable_number

def IS_inspect_document_contents():
    __set_document_name()
    #document = Document(document_name)
    global document
    __find_all_document_tables()
    __find_all_document_paragraphs()


def IS_inspect_class_description():
    result = True
    IS_antal_brister_klassbeskrivning = 0

    ############################################################
    ### DOESN'T WORK!. Compare to document paragraph instead ###
    ############################################################
    for paragraph in range(0, len(class_paragraphs_description)):
        if class_paragraphs_description[paragraph][0:2] == classes_paragraph_level_1+".":
            paragraph_number = class_paragraphs_description[paragraph][0:4].strip()
            if any(x in paragraph_number for x in class_paragraphs_number) == False:
                result = False
                #globals.IS_antal_brister_klassbeskrivning += 1
                IS_antal_brister_klassbeskrivning += 1

    if result == True:
        #write_output("<b>Resultat:</b> alla infomodellklasser har beskrivning")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser har beskrivning")
    else:
        #write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar beskrivning")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar beskrivning")
    return IS_antal_brister_klassbeskrivning

### 2do: generalisera funktionen ###
def IS_inspect_attribute_case():
    all_attributes_approved = True
    IS_antal_brister_attributnamn = 0
    for index in range(len(infomodel_table_indexes)):
        table_num = index + infomodel_table_indexes[0]
        table = document.tables[table_num]
        for row in table.rows[1:]:
            if row.cells[0].text[0:1] != row.cells[0].text[0:1].lower():
                table_title = IS_get_infomodel_classname_from_table_number(table_num, True)
                write_detail_box_content(globals.HTML_3_SPACES + "Infomodellklass " +
                        table_title + ": har fel skiftläge på attributnamn" +
                        ". Attribut: " + row.cells[0].text.strip())
                all_attributes_approved = False
                #globals.IS_antal_brister_attributnamn += 1
                IS_antal_brister_attributnamn += 1
    if all_attributes_approved == True:
        write_detail_box_content("<b>Resultat:</b> alla infomodellklassers alla attributnamn har liten begynnelsebokstav")
    else:
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser har attributnamn med stor begynnelsebokstav")
    return IS_antal_brister_attributnamn


def IS_inspect_attribute_multiplicity():
    all_results = True
    IS_antal_brister_multiplicitet = 0
    for index in range(len(infomodel_table_indexes)):
        table_num = infomodel_table_indexes[index]
        table_num = index + infomodel_table_indexes[0]
        result = __infomodel_classes_verify_multiplicity(table_num)
        # 2do
        # Insert code from __infomodel_classes_verify_multiplicity
        # 2do
        if result == False:
            tbl_no = index + infomodel_table_indexes[0]
            table_title = IS_get_infomodel_classname_from_table_number(tbl_no, True)
            #write_output(globals.HTML_3_SPACES + "Infomodellklass " + table_title + ": saknar multiplicitet")
            write_detail_box_content(globals.HTML_3_SPACES + "Infomodellklass " + table_title + ": saknar multiplicitet")
            IS_antal_brister_multiplicitet += 1
            all_results = False
    if all_results == True:
        #write_output("<b>Resultat:</b> alla infomodellklasser har multiplicitet i datatypskolumnen")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser har multiplicitet i datatypskolumnen")
    else:
        #write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar multiplicitet i datatypskolumnen")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar multiplicitet i datatypskolumnen")

    return IS_antal_brister_multiplicitet

### Data Types ###
def IS_inspect_usage_of_defined_datatypes():
    all_results = True
    IS_antal_brister_datatyper = 0
    table = document.tables[datatype_definitions_table[0]]
    defined_datatypes = []
    for i in range(1,len(table.rows)):
        defined_datatypes.append(table.cell(i, 0).text.strip().lower())

    for index in range(len(infomodel_table_indexes)):
        table_number = infomodel_table_indexes[index]
        table = document.tables[table_number]
        table_column_index = __get_table_column_index(table_number, "datatyp")
        if table_column_index == 0:
            table_column_index = __get_table_column_index(table_number, "data-typ")

        for i in range(1, len(table.rows)):
            if any(x in table.cell(i, table_column_index).text.strip().lower() for x in defined_datatypes) == False:
                table_title = __get_title_by_table_number(index)
                #write_output(globals.HTML_3_SPACES + "Datatypen är ej definierad!  Tabell: " +
                #        table_title + ". Attribut: " + table.cell(i, 0).text + ". Datatyp: " +
                #        table.cell(i, table_column_index).text.strip().lower())
                write_detail_box_content(globals.HTML_3_SPACES + "Datatypen är ej definierad!  Tabell: " +
                        table_title + ". Attribut: " + table.cell(i, 0).text + ". Datatyp: " +
                        table.cell(i, table_column_index).text.strip().lower())
                IS_antal_brister_datatyper += 1
                all_results = False
    if all_results == True:
        #write_output("<b>Resultat:</b> alla infomodellklasser använder definierade datatyper")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser använder definierade datatyper")
    else:
        #write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar eller använder icke definierad datatyp")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar eller använder icke definierad datatyp")
    return IS_antal_brister_datatyper

def IS_inspect_usage_of_reference_infomodel():
    all_results = True
    IS_antal_brister_referensinfomodell = 0
    table = document.tables[datatype_definitions_table[0]]

    for index in range(len(infomodel_table_indexes)):
        table_num = infomodel_table_indexes[index]
        table = document.tables[table_num]
        for i in range(1, len(table.rows)):
            if table.cell(i, 1).text.strip() == "":
                tbl_no = index+infomodel_table_indexes[0]
                table_title = IS_get_infomodel_classname_from_table_number(tbl_no, True)
                #write_output(globals.HTML_3_SPACES + "Mappning saknas till RIM!  Tabell: " + table_title + ". Attribut: " + table.cell(i, 0).text)
                write_detail_box_content(globals.HTML_3_SPACES + "Mappning saknas till RIM!  Tabell: " + table_title + ". Attribut: " + table.cell(i, 0).text)
                IS_antal_brister_referensinfomodell += 1
                all_results = False
    if all_results == True:
        #write_output("<b>Resultat:</b> alla infomodellklasser är mappade till referensinformationsmodellen")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser är mappade till referensinformationsmodellen")
    else:
        #write_output("<b>Resultat:</b> en eller flera infomodellklasser saknar mappning till referensinformationsmodellen")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser saknar mappning till referensinformationsmodellen")
    return IS_antal_brister_referensinfomodell


def IS_get_infomodel_classname_from_table_number(table_number, include_level):
    global infomodel_classes_list
    result_classtitle = ""
    for obj in infomodel_classes_list:
        if obj.classtable_number == table_number:
            if include_level == True:
                result_classtitle = obj.document_level + " " + obj.classtitle
            else:
                result_classtitle = obj.classtitle
            break
    return result_classtitle

def __set_document_name():
    global document_name
    """os.chdir(globals.document_path)
    for word_document in glob.glob("IS_*.doc*"):
        document_name = r""+globals.document_path+"/"+word_document"""
    document_name = IS_document_name

### Find all tables with infomodel class information ###
def __find_all_document_tables():
    global all_tables
    global datatype_definitions_table
    global document
    ##global begreppslista
    global infomodel_table_indexes
    #document = Document(document_name)
    all_tables = document.tables
    ##begreppslista = []
    datatype_definitions_table = []
    infomodel_table_indexes = []
    index = 0
    for table in all_tables:
        if table.cell(0,0).text.strip().lower() == "attribut":
            infomodel_table_indexes.append(index)
        if table.cell(0, 0).text.strip().lower() == "förkortning":
            datatype_definitions_table.append(index)
        elif table.cell(0,0).text.strip().lower() == "benämning":
            datatype_definitions_table.append(index)
        index += 1

### Find and save a list of the paragraphs ###
def __find_all_document_paragraphs():
    level_from_style_name = {f'Rubrik {i} Nr': i for i in range(10)}
    current_levels = [0] * 10

    global all_paragraphs
    global classes_paragraph_level_1
    global class_paragraphs_number
    global class_paragraphs_title
    global class_paragraphs_description
    all_paragraphs = []
    class_paragraphs_number = []
    class_paragraphs_title = []
    class_paragraphs_description = []
    classes_paragraph_level_1 = ""

    for paragraph in document.paragraphs:
        if paragraph.style.name not in level_from_style_name:
            all_paragraphs.append(paragraph.text)
        else:
            level = level_from_style_name[paragraph.style.name]
            current_levels[level] += 1
            #
            class_paragraphs_description.append(__format_levels(current_levels) + ' ' + paragraph.text)
            #
            for level in range(level + 1, 10):
                current_levels[level] = 0
            if paragraph.text[0:20].strip().lower() == "klasser och attribut":
                classes_paragraph_level_1 = __format_levels(current_levels)

            # Bug: missing paragraph descriotion => doesn't append the paragraph
            if __format_levels(current_levels)[0] == classes_paragraph_level_1:
                if __format_levels(current_levels)[0:2] == classes_paragraph_level_1 + ".":
                    class_paragraphs_title.append(__format_levels(current_levels) + ' ' + paragraph.text)
                    class_paragraphs_number.append(__format_levels(current_levels))

            all_paragraphs.append(__format_levels(current_levels) + ' ' + paragraph.text)

def __get_table_column_index(table_number, searchphrase):
    table = document.tables[table_number]
    column_count = len(table.row_cells(0))
    column_index = 0
    for i in range(0,column_count):
        if searchphrase.strip().lower() == "multiplicitet":
            if "multiplicitet" in table.cell(0, i).text.lower() or "multi-plicitet" in table.cell(0, i).text.lower():
                column_index = i
                break
        else:
            if searchphrase.strip().lower() in table.cell(0, i).text.lower():
                column_index = i
                break
    return column_index

def __get_title_by_table_number(table_number):
    table_title = ""

    ###############################################
    ### 2do: get paragraph title for this table ###
    tbl_no = table_number + infomodel_table_indexes[0]
    table_title = IS_get_infomodel_classname_from_table_number(tbl_no, True)
    if table_title.strip() == "":
        table_title = str(table_number)
    ###############################################

    return table_title

def __infomodel_classes_verify_multiplicity(table_number):
    result = True
    table = document.tables[table_number]
    allowed_multiplicities = ['0', '1', '*']
    table_column_index = __get_table_column_index(table_number, "multiplicitet")
    if table_column_index == 0:
        table_column_index = __get_table_column_index(table_number, "datatyp")
    if table_column_index == 0:
        table_column_index = __get_table_column_index(table_number, "data-typ")

    for i in range(1,len(table.rows)):
        if any(cell_multiplicity in table.cell(i, table_column_index).text for cell_multiplicity in allowed_multiplicities) == False:
            tbl_no = table_number #+ infomodel_table_indexes[0]
            table_title = IS_get_infomodel_classname_from_table_number(tbl_no, True)
            #write_output(globals.HTML_3_SPACES + globals.HTML_3_SPACES + "Klass: " + table_title + " saknar multiplicitet för: " + table.cell(i, 0).text)
            #write_detail_box_content(globals.HTML_3_SPACES + globals.HTML_3_SPACES + "Klass: " + table_title + " saknar multiplicitet för: " + table.cell(i, 0).text)
            result = False
    return result

def __inspect_revision_history():
    table = document.tables[1]
    # 2do: Display version number from page 1 in the document

    if table.cell(0,0).text == "Revisionshistorik mall":
        table = document.tables[2]

    for i, row in enumerate(table.rows):
        text = tuple(cell.text for cell in row.cells)

    if str(table.cell(i, 0).text) != globals.tag:
        #write_output("OBS! Revisionshistoriken behöver uppdateras. (hittade: "+str(table.cell(i, 0).text)+" men förväntade: "+globals.tag+")")
        write_detail_box_content("OBS! Revisionshistoriken behöver uppdateras. (hittade: "+str(table.cell(i, 0).text)+" men förväntade: "+globals.tag+")")
    else:
        #write_output("Revisionshistoriken är uppdaterad för denna version av domänen")
        write_detail_box_content("Revisionshistoriken är uppdaterad för denna version av domänen")
    #write_output("Revisionshistorikens sista rad: " + str(text))
    write_detail_box_content("Revisionshistorikens sista rad: " + str(text))

def __format_levels(current_level):
    levels = [str(level) for level in current_level if level != 0]
    return '.'.join(levels)  # Customize your format here

def __inspect_classname_format():
    result = True

    ### 2do: check that the class name begins with UPPERCASE ###
    #write_output("<br>Kontroll att infomodellklassernas rubriker är i alfabetisk ordning och börjar med stor bokstav")
    write_detail_box_content("<br>Kontroll att infomodellklassernas rubriker är i alfabetisk ordning och börjar med stor bokstav")
    for paragraph in range(0, len(class_paragraphs_title)):
        #write_output(globals.HTML_3_SPACES + class_paragraphs_title[paragraph])
        write_detail_box_content(globals.HTML_3_SPACES + class_paragraphs_title[paragraph])
    #write_output("För närvarande sker kontrollen manuellt, med ovanstående listning som underlag")
    write_detail_box_content("För närvarande sker kontrollen manuellt, med ovanstående listning som underlag")

    return result

def __infomodel_classes_verify_attribute_case(table_number):
    table = document.tables[table_number]
    result = True
    for i in range(1,len(table.rows)):
        if table.cell(i, 0).text[0:1].islower() == False:
            #write_output(globals.HTML_3_SPACES + "Fel skiftläge för första bokstaven i: " + table.cell(i, 0).text)
            write_detail_box_content(globals.HTML_3_SPACES + "Fel skiftläge för första bokstaven i: " + table.cell(i, 0).text)
            result = False
    return result

def __inspect_attribut_name_case():
    all_results = True
    #write_output("\nKontroll att attributnamn inleds med liten bokstav i infomodellklassernas tabeller")
    write_detail_box_content("\nKontroll att attributnamn inleds med liten bokstav i infomodellklassernas tabeller")
    for index in range(len(infomodel_table_indexes)):
        table_num = infomodel_table_indexes[index]
        result = __infomodel_classes_verify_attribute_case(table_num)
        # 2do
        # Insert code from __infomodel_classes_verify_attribute_case
        # 2do
        if result == False:
            #write_output(globals.HTML_3_SPACES + "Infomodellklass " + str(index+1) + ": innehåller attributnamn med fel skiftläge för första bokstaven\n")
            write_detail_box_content(globals.HTML_3_SPACES + "Infomodellklass " + str(index+1) + ": innehåller attributnamn med fel skiftläge för första bokstaven\n")
            all_results = False
    if all_results == True:
        #write_output("<b>Resultat:</b> alla infomodellklasser är OK avseende skiftläge på attributnamnens första bokstav")
        write_detail_box_content("<b>Resultat:</b> alla infomodellklasser är OK avseende skiftläge på attributnamnens första bokstav")
    else:
        #write_output("<b>Resultat:</b> en eller flera infomodellklasser har fel skiftläge för första bokstaven i attribut")
        write_detail_box_content("<b>Resultat:</b> en eller flera infomodellklasser har fel skiftläge för första bokstaven i attribut")

def IS_init_infomodel_classes_list():
    global document
    global document_name
    __set_document_name()
    #print("document_name:",document_name)
    #print("globals.docx_IS_document",globals.docx_IS_document)
    #document = Document(document_name)
    #document = Document(globals.IS_document_name)
    document = globals.docx_IS_document
    #print("document",document)
    __find_all_document_paragraphs()
    __find_all_document_tables()

    global infomodel_classes_list
    infomodel_classes_list = []
    searched_paragraph_level = 2    #2do: replace with correct value
    append_paragraph_title = ""
    paragraph_level = ""
    table_number = infomodel_table_indexes[0]

    for block in __iter_block_items(document, searched_paragraph_level):
        append_table_number = 0
        if isinstance(block, Paragraph):
            new_paragraph_level = DOCX_display_document_contents.DOCX_document_structure_get_levelvalue(block.text)
            if new_paragraph_level != globals.NOT_FOUND:
                paragraph_level = new_paragraph_level
            if len(paragraph_level) == 1:
                continue
            if paragraph_level != "" and paragraph_level[0] == classes_paragraph_level_1[0]:
                if DOCX_display_document_contents.DOCX_document_structure_get_levelvalue(block.text) != globals.NOT_FOUND:
                    append_paragraph_title = block.text
        elif isinstance(block, Table):
            if paragraph_level != "":
                if paragraph_level[0] == classes_paragraph_level_1[0]:
                    append_table_number = table_number
                    table_number += 1

        if append_paragraph_title != "" and append_table_number != 0:
            #print("\tappend:", append_paragraph_title, paragraph_level, append_table_number)
            append_paragraph_title = append_paragraph_title.replace("\n"," ")
            infomodel_classes_list.append(infomodel_classes(append_paragraph_title, paragraph_level, append_table_number))
            append_paragraph_title == ""

        #for obj in infomodel_classes_list:
        #    print("\t", obj.document_level, obj.classtitle, obj.classtable_number, sep=' ')
        #print("Infomodellklass för tabell 12:", __get_infomodel_classname_from_table_number(12,False))

    #pet: test
    IS_init_all_tables_list()
    #pet: test


def IS_init_all_tables_list():
    global document
    #__set_document_name()
    #document = Document(document_name)
    #__find_all_document_paragraphs()
    #__find_all_document_tables()

    global all_tables_list
    all_tables_list = []
    searched_paragraph_level = 1    #2do: replace with correct value
    append_paragraph_title = ""
    paragraph_level = ""
    table_number = 0    #infomodel_table_indexes[0]

    for block in __iter_block_items(document, searched_paragraph_level):
        append_table_number = 0
        if isinstance(block, Paragraph):
            new_paragraph_level = DOCX_display_document_contents.DOCX_document_structure_get_levelvalue(block.text)
            if new_paragraph_level != globals.NOT_FOUND:
                paragraph_level = new_paragraph_level
            if len(paragraph_level) == 1:
                continue
            if paragraph_level != "": #and paragraph_level[0] == classes_paragraph_level_1[0]:
                if DOCX_display_document_contents.DOCX_document_structure_get_levelvalue(block.text) != globals.NOT_FOUND:
                    append_paragraph_title = block.text
        elif isinstance(block, Table):
            if paragraph_level != "":
                #if paragraph_level[0] == classes_paragraph_level_1[0]:
                #    append_table_number = table_number
                #    table_number += 1
                append_table_number = table_number
                table_number += 1

        if append_paragraph_title != "" and append_table_number != 0:
            #print("\tappend:", append_paragraph_title, paragraph_level, append_table_number)
            append_paragraph_title = append_paragraph_title.replace("\n"," ")
            all_tables_list.append(infomodel_classes(append_paragraph_title, paragraph_level, append_table_number))
            append_paragraph_title == ""
    #print("for obj in all_tables_list")
    #for obj in all_tables_list:
    #    print("\t", obj.document_level, obj.classtitle, obj.classtable_number, sep=' ')
    #print("Infomodellklass för tabell 12:", __get_infomodel_classname_from_table_number(12,False))


def __iter_block_items(parent,searched_paragraph_level):
    if isinstance(parent, docx.document.Document):   #_Document
        parent_elm = parent.element.body

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)



"""def __get_infomodel_table_number_from_classname(classtitle):
    global infomodel_classes_list
    result_table_number = ""
    for obj in infomodel_classes_list:
        classtitle_without_level = obj.classtitle.strip().lower()
        if classtitle.strip().lower() == classtitle_without_level:
            result_table_number = obj.classtable_number
            break
    return result_table_number"""

"""def IS_get_tableno_for_first_title_cell(title):
    global document
    global all_tables   # equals to document.tables
    table_number = 0
    index = 0
    for table in all_tables:
        if table.cell(0,0).text.strip().lower() == title:
            table_number = index
            break
        index += 1
    return table_number"""

"""def __inspect_class_order():
    ### 2do ###
    result = True"""

"""def __document_structure_get_key(searched_value):
    for key, value in document_structure_dict.items():
        if searched_value == value:
            return key
    return NOT_FOUND"""
