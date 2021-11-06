import datetime
from docx.document import Document as _Document
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
import globals
import granskning_AB
import granskning_IS
import granskning_TKB
from utilities import *


INITIAL_NEWLINE = True
NO_INITIAL_NEWLINE = False
NO_TABLES = False
NO_TEXT = False
NO_TITLE = False
NOT_FOUND = "Not found"
STYLE_FAMILY_HEADING = "Heading"
STYLE_FAMILY_RUBRIK = "Rubrik"
STYLE_FAMILY_SUBTLE_EMPHASIS = "Subtle Emphasis"
TABLES = True
TEXT = True
TITLE = True


def __style_family(document_search_phrase):
    if "AB_*" in document_search_phrase:
        return STYLE_FAMILY_RUBRIK
    elif "IS_*" in document_search_phrase:
        return STYLE_FAMILY_RUBRIK
    else:
        return STYLE_FAMILY_HEADING


def DOCX_prepare_inspection(document_search_phrase):
    """
    Anropar metoder som förbereder granskning av ett Worddokument
    """
    __set_document_name(document_search_phrase)

    __set_document_name(document_search_phrase)
    __document_structure_2_dict(__style_family(document_search_phrase))

    DOCX_init_dict_paragraph_title_and_tableno(document)

def DOCX_inspect_revision_history(docx_document, table_num):
    """
    Kollar att dokumentets tabell med revisionshistorik har en rad för aktuell tag.
    """
    table = document.tables[table_num]
    antal_brister_revisionshistorik = 0
    # 2do: Display version number from page 1 in the document

    if table.cell(0,0).text == "Revisionshistorik mall":
        table = document.tables[table_num+1]

    for i, row in enumerate(table.rows):
        text = tuple(cell.text for cell in row.cells)

    if str(table.cell(i, 0).text) != globals.tag:
        write_detail_box_content("<b>Resultat:</b> Revisionshistoriken behöver uppdateras. (hittade: "+str(table.cell(i, 0).text)+" men förväntade: "+globals.tag+")")
        if docx_document == globals.IS:
            antal_brister_revisionshistorik = 1
        elif docx_document == globals.TKB:
            antal_brister_revisionshistorik = 1
        elif docx_document == globals.AB:
            antal_brister_revisionshistorik = 1
    else:
        write_detail_box_content("<b>Resultat:</b> Revisionshistoriken är uppdaterad för denna version av domänen")
    write_detail_box_content("Revisionshistorikens sista rad: " + str(text))
    #print(document.core_properties.author)

    # 2do: kolla om dokumentet har eftersökta custom properties
    # Custom properties finns i custom.xml i Wordfilen
    """try:
        print(document.custom_properties)
        #print("Version",document.CustomDocumentProperties('Version').value)
    except AttributeError:
        print(globals.docx_document, "has no custom properties")    # Hittills kommer alla anrop hit
    #for paragraph in document.paragraphs:
    #    print(globals.docx_document,"paragraph",paragraph.text)"""

    return antal_brister_revisionshistorik

def DOCX_inspect_revision_history_new(docx_document, table):
    """
    Kollar att dokumentets tabell med revisionshistorik har en rad för aktuell tag.
    Optimerad för att undvika timeout då revisionshistoriktabellen har (alltför) många rader
    """
    antal_brister_revisionshistorik = 0

    for i, row in enumerate(table.rows):
        text = tuple(cell.text for cell in row.cells)

    if str(table.cell(i, 0).text) != globals.tag:
        write_detail_box_content("<b>Resultat:</b> Revisionshistoriken behöver uppdateras. (hittade: "+str(table.cell(i, 0).text)+" men förväntade: "+globals.tag+")")
        antal_brister_revisionshistorik = 1
    else:
        write_detail_box_content("<b>Resultat:</b> Revisionshistoriken är uppdaterad för denna version av domänen")
    write_detail_box_content("Revisionshistorikens sista rad: " + str(text))

    return antal_brister_revisionshistorik


def DOCX_display_paragraph_text_and_tables(searched_paragraph_title, display_paragraph_title, display_initial_newline, display_keylevel_text, display_tables):
    """
    Skriver ut titlar och innehåll i paragrafer och tabeller i enlighet med angivna parametrar.

    Utskrift sker via metoden: write_output

    Förbättringsförslag: lägg till en inparameter för att kunna välja om searched_paragraph_title ska matchas exakt eller ingå i hittad titel
    """
    searched_paragraph_level = DOCX_document_structure_get_levelvalue(searched_paragraph_title)
    searched_paragraph_found = False
    paragraph_or_table_found = False

    if display_paragraph_title == True and display_tables == False:
        if display_initial_newline == True:
            write_detail_box_content("<br>")
        paragraph_displayed = __display_paragraph_text_by_paragraph_level(searched_paragraph_level, display_keylevel_text)
        if paragraph_displayed == True:
            paragraph_or_table_found = True
    else:
        for block in __document_block_items(document):      #__iter_block_items(document,searched_paragraph_level)
            if isinstance(block, Paragraph):
                this_paragraph_title = block.text.strip().lower()
                if this_paragraph_title == searched_paragraph_title.strip().lower():
                    searched_paragraph_found = True
                    paragraph_or_table_found = True
                    if display_paragraph_title == True:
                        __display_paragraph_text_by_paragraph_level(searched_paragraph_level, display_keylevel_text)

            elif isinstance(block, Table):
                if searched_paragraph_found == True:
                    if display_tables == True:
                        if display_paragraph_title == False:
                            write_detail_box_content("<br>")
                        __document_table_print_html_table(block)
                        paragraph_or_table_found = True
                    searched_paragraph_found = False     # Bug: supports only one table per paragraph

    return paragraph_or_table_found

def DOCX_list_searched_paragraph_titles_wrong_case(searched_paragraph_title, delimiter, searched_case):
    searched_paragraph_level = DOCX_document_structure_get_levelvalue(searched_paragraph_title)
    paragraph_title_list = []

    for block in __document_block_items(document):   # )__iter_block_items(document, searched_paragraph_level)
        if isinstance(block, Paragraph):
            this_paragraph_title = block.text.strip()
            if this_paragraph_title.lower() == searched_paragraph_title.strip().lower():

                previous_key = ""
                count = 0
                for key, value in document_paragraph_index_dict.items():
                    if key[0:len(searched_paragraph_level)] == searched_paragraph_level:
                        key_level_length = key.find(" ")
                        if len(key.strip()) > key_level_length:
                            this_key_level = key.strip()[0:key_level_length]
                            if this_key_level != previous_key:
                                key = key.replace("\n", " ")
                                key_extract = key[key.find(delimiter)+len(delimiter):]
                                if key[key.find(" "):].lower().strip() != searched_paragraph_title.lower().strip():
                                    if key[key.find(delimiter):].lower() != searched_paragraph_title.lower():
                                        if searched_case == globals.UPPER_CASE:
                                            if key_extract[0] != key_extract[0].upper():
                                                paragraph_title_list.append(key)
                                                count += 1
                                        else:
                                            if key_extract[0] != key_extract[0].lower():
                                                paragraph_title_list.append(key)
                                                count += 1
                            previous_key = key.strip()[0:key_level_length]
    return paragraph_title_list, count

def DOCX_inspect_reference_links(table_num):
    """
    Kollar om länkarna i dokumentets referenstabell fungerar eller ej.
    """
    antal_brister_referenslänkar = 0
    links_excist = False
    links = extract_urls_from_table(document, table_num)
    if len(links) == 0:
        write_detail_box_content("Det finns inga länkar i referenstabellen. Obs att det ändå kan förekomma länkar med annat format (text istället för hyperlänk).")
    else:
        links_excist = True
    for link in links:
        begin = datetime.datetime.now()
        status_code = verify_url_exists(link)
        end = datetime.datetime.now()
        diff = end-begin
        if status_code == 400:
            write_detail_box_content("<b>Länken är felaktig eller kan inte tolkas!</b> (statuskod: " + str(status_code) + ") för: " + link)
            antal_brister_referenslänkar += 1
        elif status_code < 404:
            write_detail_box_content(globals.HTML_3_SPACES + "<b>OK</b> (statuskod: " + str(status_code) + ") för: <a href='" + link + "' target = '_blank'>" + link + "</a>")
        else:
            antal_brister_referenslänkar += 1
            write_detail_box_content(globals.HTML_3_SPACES + "<b>Sidan saknas!</b> (statuskod: " + str(status_code) + ") för: " + link)

    return links_excist, antal_brister_referenslänkar

#def DOCX_display_paragragh_title(searched_title_name):
    """
    Söker efter angiven paragraf i lagrad dokumentstruktur. Skriver ut paragrafens titel.

    Returnerar: True om sökt paragraf hittades och False om paragrafen inte hittades
    """
    """result = True
    searched_paragraph_level = DOCX_document_structure_get_exact_levelvalue(searched_title_name)
    if searched_paragraph_level != "":
        #write_output("OK. (" + searched_title_name + ") avsnitt " + searched_paragraph_level + " i TKB")
        write_output("OK. Dokument-rubrik (" + searched_paragraph_level + "):  \t" + searched_title_name)
    else:
        write_output("FEL! " + searched_title_name + " verkar inte vara beskrivet i dokumentet!")
        result = False
    return result"""


def __set_document_name(search_phrase):
    """
    Sätter den globala variabeln 'document' till namnet på angivet dokument
    """
    global document

    if "IS_*" in search_phrase:
        document = globals.docx_IS_document
    elif "TKB_*" in search_phrase:
        document = globals.docx_TKB_document
    elif "AB_*" in search_phrase:
        document = globals.docx_AB_document


def __document_structure_2_dict(style_family):
    """
    Lagrar dokumentet struktur i ett globalt dictionary.
        Key = rubriktext
        Value = rubriknivå (exempelvis 2.1)

    Lagrar även index till dokumentets paragrafer i ett globalt dictionary.
        Key = rubriknivå (exempelvis 2.1) + " " + rubriktext
        Value = indexnummer

    """
    if style_family == STYLE_FAMILY_HEADING:
        level_from_style_name = {f'Heading {i}': i for i in range(10)}
    elif style_family == STYLE_FAMILY_RUBRIK:
        level_from_style_name = {f'Rubrik {i} Nr': i for i in range(10)}
    elif style_family == STYLE_FAMILY_SUBTLE_EMPHASIS:
        #2do: try other alternatives to make this style work as title
        level_from_style_name = {i for i in range(2)}
    current_levels = [0] * 10
    global document_structure_dict  #key = kapitelnamn, value = kapitelnummer
    document_structure_dict = {}
    global document_paragraph_index_dict
    document_paragraph_index_dict = {}
    global paragraph_title_tableno_dict
    paragraph_title_tableno_dict = {}

    index = 1
    for paragraph in document.paragraphs:
        if paragraph.style.name in level_from_style_name:
            level = level_from_style_name[paragraph.style.name]
            current_levels[level] += 1
            for level in range(level + 1, 10):
                current_levels[level] = 0
            document_structure_dict[paragraph.text.strip().lower()] = __format_levels(current_levels)
        document_paragraph_index_dict[__format_levels(current_levels) + " " + paragraph.text] = index
        index +=1


def DOCX_document_structure_get_levelvalue(searched_key):
    """
    Söker efter angivet rubrikvärde i dictionaryt med dokumentstruktur.

    Returnerar: Om rubrikvärdet hittades så returneras dess nyckel (rubrikens titel), annars returneras NOT_FOUND
    """
    for key, value in document_structure_dict.items():
        if searched_key.strip().lower() == key:
            return value
    return NOT_FOUND


def __display_paragraph_text_by_paragraph_level(searched_paragraph_level,display_keylevel_text):
    """
    Hämtar paragraftext från dokumentstruktur-dictionaryt med rubriknivå som nyckel, och visar den funna texten
    """
    paragraph_displayed = False
    previous_key = ""
    for key, value in document_paragraph_index_dict.items():
        if key[0:len(searched_paragraph_level)] == searched_paragraph_level:
            key_level_length = key.find(" ")
            if len(key.strip()) > key_level_length:
                this_key_level = key.strip()[0:key_level_length]
                if this_key_level == previous_key:
                    if display_keylevel_text == True:
                        write_detail_box_content(globals.HTML_3_SPACES + key.strip()[key_level_length+1:])
                        paragraph_displayed = True
                else:
                    key = key.replace("\n"," ")
                    write_detail_box_content(globals.HTML_3_SPACES + key)
                    paragraph_displayed = True
                previous_key = key.strip()[0:key_level_length]
    return paragraph_displayed


def __format_levels(current_level):
    """
    Formaterar angiven rubriknivå. 2do: beskriv detta bättre...

    Returnerar: formaterad rubriknivå
    """
    levels = [str(level) for level in current_level if level != 0]
    return '.'.join(levels)

def DOCX_document_structure_get_key(searched_value):
    for key, value in document_structure_dict.items():
        if searched_value == value:
            return key
    return NOT_FOUND


def __document_table_print_html_table(table):
    html_table = "<style> table, th, td { border:1px solid gray; empty-cells: show; } </style>"
    html_table += "<table>"
    #html_table += "<caption>"+title+"</caption>"
    row_number = 0
    for row in table.rows:
        row_number += 1
        html_table += "<tr>"
        for cell in row.cells:
            if row_number == 1:
                html_table += "<th>" + cell.text.strip() + "</th>"
            else:
                html_table += "<td>" + cell.text.strip() + "</td>"
        html_table += "</tr>"
    html_table += "</table>"
    write_detail_box_content(html_table)


def DOCX_empty_table_cells_exists(table_number, display_result, display_type):
    result = False
    if globals.docx_document == globals.IS:
        granskning_IS.IS_antal_brister_tomma_tabellceller = 0
    elif globals.docx_document == globals.TKB:
        granskning_TKB.TKB_antal_brister_tomma_tabellceller = 0
    elif globals.docx_document == globals.AB:
        granskning_AB.AB_antal_brister_tomma_tabellceller = 0
    antal_brister_tomma_tabellceller = 0

    html_table = ""
    if display_type == globals.DISPLAY_TYPE_TABLE:
        html_table += "<table>"

    table = document.tables[table_number]

    row_number = 0
    for i,row in enumerate(table.rows):
        row_number += 1
        column_count = len(table.row_cells(0))
        cells_missing_content = ""
        cell_contents_html = ""
        table_title = ""
        celltext = tuple(cell.text for cell in row.cells)
        cell_number = 0
        for element in celltext:
            cell_has_contents = False
            if row_number == 1:
                html_table += "<th>" + element + "</th>"
            cell_number += 1
            if element != "":
                cell_has_contents = True
                if cell_contents_html == "":
                    cell_contents_html += "<tr>"
                cell_contents_html += "<td>" + element + "&nbsp;</td>"
            else:
                if cell_contents_html == "":
                    cell_contents_html += "<tr>"
                cell_contents_html += "<td>" + element + "&nbsp;</td>"
                # Hyperlinks are represented as empty in tuples. Check for hyperlink in paragraph XML
                for paragraph in table.cell(row_number-1,cell_number-1).paragraphs:
                    xml_str = str(paragraph.paragraph_format.element.xml)
                    if "<w:t>" in xml_str or "<w:hyperlink" in xml_str or 'w:val="Hyperlink"' in xml_str:
                        cell_has_contents = True
            if cell_has_contents == False:
                result = True
                antal_brister_tomma_tabellceller += 1
                if globals.docx_document == globals.IS:
                    table_title = granskning_IS.IS_get_infomodel_classname_from_table_number(table_number, True)
                elif globals.docx_document == globals.TKB:
                    table_title = "TKB-tabell nummer " + str(table_number)
                elif globals.docx_document == globals.AB:
                    table_title = "AB-tabell nummer " + str(table_number)
                if cells_missing_content == "":
                    cells_missing_content += str(cell_number+1)
                else:
                    cells_missing_content += ", " + str(cell_number+1)   #column
        cell_contents_html += "</tr>"
        if cells_missing_content != "":
            if display_type == globals.DISPLAY_TYPE_TEXT:
                write_detail_box_content(
                    globals.HTML_3_SPACES + "Tabellceller utan innehåll!  Tabell: " + table_title + ", Rad: " + str(
                        row) + ", Kolumn: " + cells_missing_content)
            elif display_type == globals.DISPLAY_TYPE_TABLE and result == True:
                html_table += cell_contents_html

    if display_result == True:
        if result == True:
            if display_type == globals.DISPLAY_TYPE_TABLE:
                html_table += "</table>"
                write_detail_box_content(html_table)
            write_detail_box_content("<b>Resultat:</b> det finns granskade tabell(er) med en eller flera celler utan innehåll")
        else:
            write_detail_box_content("<b>Resultat:</b> alla granskade celler har innehåll")

    return result, antal_brister_tomma_tabellceller


### Funktionen borde kunna arbetas bort, den används bara som fallback i IS då bepreppstabell inte hittas ###
def DOCX_get_tableno_for_first_column_title(title, all_tables):
    table_number = 0
    index = 0
    for table in all_tables:
        if table.cell(0,0).text.strip().lower() == title.strip().lower():
            table_number = index
            break
        index += 1

    return table_number


def DOCX_init_dict_paragraph_title_and_tableno(document):
    """
        Tabeller som har likadana rubriker, exempelvis "fältregler", lagrar bara en post i dictionaryt i funktionens grundutförande
            Detta är åtgärdat genom att konkatenera nyckeln med tabellnumret
        TKB- och AB-dokumenten innehåller oftast en tabell på sida 1. Tabellen anger domän, version och datum, men har ingen rubrik
            Dessa tabeller sparas inte i dictionaryt pga att de saknar rubrik

        Key: rubrik
        Value: tabellnummer
    """

    global paragraph_title_tableno_dict

    paragraph_text = ""
    table_index = 0
    for block in __document_block_items(document):
        if isinstance(block, Paragraph):
            if block.text.strip() != "" and block.style.name.lower() != "normal"  and block.style.name.lower() != "body text":
                paragraph_text = block.text.strip().lower()
        elif isinstance(block, Table):
            if block.table.cell(0, 0).text.strip() != "":
                table_index += 1
                if paragraph_text == "":
                    paragraph_text = "Tabell " + str(table_index) + ": " + block.table.cell(0, 0).text
                if paragraph_text in paragraph_title_tableno_dict:
                    paragraph_title_tableno_dict[paragraph_text + "." + str(table_index)] = table_index
                else:
                    paragraph_title_tableno_dict[paragraph_text] = table_index
                paragraph_text = ""

    ### Stöd vid utveckling och felsökning ###
    """print("\n"+globals.docx_document+"\t("+globals.domain_name+", "+globals.tag+")")
    print("paragraph_title_tableno_dict:")
    for key, value in paragraph_title_tableno_dict.items():
        print("\t",value,key)"""

    return paragraph_title_tableno_dict

def __document_block_items(document):
    """
    Returnerar ett generatorobjekt med paragrafer och tabeller
    Generatorn kan användas vid iteration av dokumentets innehåll
    """
    if isinstance(document, _Document):
        document_element = document.element.body

    for child in document_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def DOCX_get_tableno_for_paragraph_title(title):
    """
        Anrop:  begreppsbeskrivning_tabell = DOCX_get_tableno_for_paragraph_title("Beskrivning av begrepp")
        Impl:   return paragraph_title_tableno_dict[title]
    """
    table_number = -1
    global paragraph_title_tableno_dict
    if title.lower() in paragraph_title_tableno_dict:
        table_number = paragraph_title_tableno_dict[title]

    return table_number
