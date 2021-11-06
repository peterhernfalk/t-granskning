"""import Document_mangagement
#from TKB_inspection import *
#from utilities import write_output, write_detail_box_content, verify_url_exists, check_if_file_exists
from utilities import *"""
import datetime

import DOCX_display_document_contents
#import html_dashboard
from DOCX_display_document_contents import *
import docx
from docx.table import *
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import *
from docx.oxml.table import *
from docx.api import Document  # noqa
import document_mangagement
###from utilities import *
import utilities

### From globals.py ###
TKB_antal_brister_referenslänkar = 0
TKB_antal_brister_revisionshistorik = 0
TKB_antal_brister_tomma_referenstabellceller = 0
TKB_antal_brister_tomma_revisionshistoriktabellceller = 0
#TKB_antal_brister_tomma_tabellceller = 0
TKB_detail_box_contents = ""
TKB_document_exists = False
TKB_document_name = ""
TKB_exists = False
TKB_meddelandemodeller_finns = False
#######################

TITLE = True
NO_TITLE = False
INITIAL_NEWLINE = True
NO_INITIAL_NEWLINE = False
TEXT = True
NO_TEXT = False
TABLES = True
NO_TABLES = False

TABLE_NUM_REVISION = 1   #Hard coded, assuming that the reference table is number 1 in the document
TABLE_NUM_REF = 2   #Hard coded, assuming that the reference table is number 2 in the document


def prepare_TKB_inspection(domain, tag, alt_document_name):
    """
    Beräknar url till TKB-dokumentet för angiven domain och tag.

    Laddar ner dokumentet till en virtuell fil som läses in i ett docx-Document.

    Anropar därefter metoden "INFO_inspect_document" som genomför granskning av dokumentet.
    """
    """
    2do: Förenkla och snygga till koden
    """
    print("TKB-init påbörjas",datetime.datetime.now().replace(microsecond=0))
    ### From globals.py ###
    global TKB_antal_brister_referenslänkar
    global TKB_antal_brister_revisionshistorik
    global TKB_antal_brister_tomma_referenstabellceller
    global TKB_antal_brister_tomma_revisionshistoriktabellceller
    global TKB_detail_box_contents
    global TKB_document_exists
    global TKB_document_name
    global TKB_exists
    global TKB_meddelandemodeller_finns
    TKB_antal_brister_referenslänkar = 0
    TKB_antal_brister_revisionshistorik = 0
    TKB_antal_brister_tomma_referenstabellceller = 0
    TKB_antal_brister_tomma_revisionshistoriktabellceller = 0
    #TKB_antal_brister_tomma_tabellceller = 0
    TKB_detail_box_contents = ""
    TKB_document_exists = False
    TKB_document_name = ""
    TKB_exists = False
    TKB_meddelandemodeller_finns = False
    #######################

    global TKB_page_link
    global TKB_document_paragraphs
    #TKB_page_link = __get_document_page_link(domain, tag, globals.TKB)
    #downloaded_TKB_page = __get_downloaded_document(TKB_page_link)
    TKB_page_link = document_mangagement.DOC_get_document_page_link(domain, tag, globals.TKB)
    downloaded_TKB_page = document_mangagement.DOC_get_downloaded_document(TKB_page_link)

    TKB_document_paragraphs = ""

    TKB_head_hash = document_mangagement.DOC_get_head_hash(downloaded_TKB_page)
    TKB_document_link = document_mangagement.DOC_get_document_link(domain, tag, globals.TKB, TKB_head_hash, alt_document_name)
    downloaded_TKB_document = document_mangagement.DOC_get_downloaded_document(TKB_document_link)
    if downloaded_TKB_document.status_code == 404:
        ###TKB_document_paragraphs = APP_text_document_not_found(globals.TKB, domain, tag)
        ###globals.granskningsresultat += "<br><br><h2>TKB</h2>" + APP_text_document_not_found(globals.TKB, domain, tag)
        docx_TKB_document = ""
        #globals.TKB_felmeddelande = APP_text_document_not_found(globals.TKB, domain, tag)
        TKB_exists = False
    else:
        globals.docx_TKB_document = document_mangagement.DOC_get_docx_document(downloaded_TKB_document)
        globals.TKB_document_exists = True
        TKB_exists = True
        ### dev test ###
        for paragraph in globals.docx_TKB_document.paragraphs:
            if paragraph.text.strip() != "":
                TKB_document_paragraphs += paragraph.text + "<br>"
        ### dev test ###

        DOCX_display_document_contents.DOCX_prepare_inspection("TKB_*.doc*")

def perform_TKB_inspection(domain, tag, alt_document_name):
    prepare_TKB_inspection(domain, tag, alt_document_name)
    print("TKB-granskning påbörjas",datetime.datetime.now().replace(microsecond=0))

    global TKB_antal_brister_referenslänkar
    global TKB_antal_brister_revisionshistorik
    global TKB_antal_brister_tomma_referenstabellceller
    global TKB_antal_brister_tomma_revisionshistoriktabellceller
    global TKB_detail_box_contents
    global TKB_document_exists
    global TKB_document_name
    global TKB_exists
    global TKB_meddelandemodeller_finns

    if TKB_exists == False:
        return


    #write_detail_box_content("<br>")
    #write_detail_box_content("<b>Krav:</b> ResultCode ska inte förekomma i läsande tjänster (kollas av RIVTA:s verifieringsscript)")
    #write_detail_box_content("<b>Krav:</b> för uppdaterande tjänster som kan returnera returkoder ska det finnas beskrivning av hur ResultCode ska hanteras")

    utilities.write_detail_box_content("<br>")
    utilities.write_detail_box_content("<b>Krav:</b> om dokumentegenskaper finns ska version och ändringsdatum stämma överens med granskad version")
    # 2do: kontrollera dokumentegenskaper avseende versionsnummer
    # 2do: kontrollera versionsnummer på dokumentets första sida
    utilities.write_detail_box_content("<b>Granskningsstöd:</b> alla interaktioner ska vara beskrivna i TKB")

    print("\tTKB: revisionshistorik, version",datetime.datetime.now().replace(microsecond=0))
    utilities.write_detail_box_content("<br>")
    utilities.write_detail_box_content("<b>Krav:</b> revisionshistoriken ska vara uppdaterad för samma version som domänen")
    utilities.write_detail_box_content("<b>Granskningsstöd:</b> om revisionshistoriken inte är uppdaterad, kontakta beställaren eller skriv en granskningskommentar")
    used_table_no = DOCX_display_document_contents.DOCX_get_tableno_for_paragraph_title("revisionshistorik")
    if used_table_no > 0:
        TKB_antal_brister_revisionshistorik = DOCX_inspect_revision_history(globals.TKB, used_table_no)
        #globals.TKB_antal_brister_revisionshistorik = DOCX_inspect_revision_history_new(globals.TKB, globals.docx_TKB_document.tables[used_table_no])
    else:
        TKB_antal_brister_revisionshistorik = DOCX_display_document_contents.DOCX_inspect_revision_history(globals.TKB,TABLE_NUM_REVISION)
        #globals.TKB_antal_brister_revisionshistorik = DOCX_inspect_revision_history_new(globals.TKB,globals.docx_TKB_document.tables[used_table_no])

    print("\tTKB: revisionshistorik, tomma celler",datetime.datetime.now().replace(microsecond=0))
    utilities.write_detail_box_content("<br>")
    utilities.write_detail_box_content("<b>Krav:</b> revisionshistorikens alla tabellceller ska ha innehåll")
    if used_table_no > 0:
        #result, TKB_antal_brister_tomma_revisionshistoriktabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(used_table_no, True, globals.DISPLAY_TYPE_TABLE)
        result, TKB_antal_brister_tomma_revisionshistoriktabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(used_table_no, True, globals.DISPLAY_TYPE_TABLE)
    else:
        ###result, TKB_antal_brister_tomma_revisionshistoriktabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(TABLE_NUM_REVISION, True, globals.DISPLAY_TYPE_TABLE)
        print("\tTKB: revisionshistorik, tomma celler, avslutad kontroll", datetime.datetime.now().replace(microsecond=0))
        result, TKB_antal_brister_tomma_revisionshistoriktabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(TABLE_NUM_REVISION, True, globals.DISPLAY_TYPE_TABLE)

    print("\tTKB: referenstabell, länkar",datetime.datetime.now().replace(microsecond=0))
    utilities.write_detail_box_content("<br>")
    utilities.write_detail_box_content("<b>Krav:</b> länkarna i referenstabellen ska fungera")
    used_table_no = DOCX_display_document_contents.DOCX_get_tableno_for_paragraph_title("referenser")
    links_excist = False
    if used_table_no > 0:
        links_excist, TKB_antal_brister_referenslänkar = DOCX_display_document_contents.DOCX_inspect_reference_links(used_table_no)
    else:
        links_excist, TKB_antal_brister_referenslänkar = DOCX_display_document_contents.DOCX_inspect_reference_links(TABLE_NUM_REF)
    if TKB_antal_brister_referenslänkar > 0:
        utilities.write_detail_box_content("<b>Resultat:</b> en eller flera länkar är felaktiga, eller kan inte tolkas korrekt av granskningsfunktionen.")
        utilities.write_detail_box_content("<b>Granskningsstöd:</b> gör manuell kontroll i dokumentet av de länkar som rapporteras som felaktiga")
    else:
        if links_excist == True:
            utilities.write_detail_box_content("<b>Resultat:</b> alla kontrollerade länkar fungerar")
        else:
            utilities.write_detail_box_content("<b>Resultat:</b> inga länkar har kontrollerats")

    utilities.write_detail_box_content("<br>")
    utilities.write_detail_box_content("<b>Krav:</b> referenstabellens alla tabellceller ska ha innehåll")
    if used_table_no > 0:
        result, TKB_antal_brister_tomma_referenstabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(used_table_no, True, globals.DISPLAY_TYPE_TABLE)
    else:
        result, TKB_antal_brister_tomma_referenstabellceller = DOCX_display_document_contents.DOCX_empty_table_cells_exists(TABLE_NUM_REF, True, globals.DISPLAY_TYPE_TABLE)

    # 2do: kontrollera om domännamnet nämns i inledningsparagrafen (det ska vara på engelska)
    # 2do: visa innehåll i inledningens underparagraf (Svenskt namn), för manuell kontroll av svenskt namn och svenskt kortnamn

    print("\tTKB: versionskontroll",datetime.datetime.now().replace(microsecond=0))
    utilities.write_detail_box_content("<br>")
    utilities.write_detail_box_content("<b>Krav:</b> versionsnumret ska vara uppdaterat för samma version som domänen")
    utilities.write_detail_box_content("<b>Krav:</b> ändringsstatus för tjänstekontrakt ska överensstämma med granskningsbeställningen")
    DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables("versionsinformation",TITLE,NO_INITIAL_NEWLINE,TEXT,NO_TABLES)
    utilities.write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående listning som underlag")

    print("\tTKB: meddelandemodeller",datetime.datetime.now().replace(microsecond=0))
    utilities.write_detail_box_content("<br>")
    utilities.write_detail_box_content("<b>Krav:</b> TKB ska innehålla ett avsnitt för meddelandemodeller")
    TKB_meddelandemodeller_finns = DOCX_display_document_contents.DOCX_display_paragraph_text_and_tables("Tjänstedomänens meddelandemodeller", TITLE, NO_INITIAL_NEWLINE, NO_TEXT, NO_TABLES)
    if TKB_meddelandemodeller_finns == False:
        utilities.write_detail_box_content("<b>Granskningsstöd:</b> inget innehåll visas, vilket kan bero på att avsnittsrubriken saknas eller är annan än den förväntade (Tjänstedomänens meddelandemodeller)")
    utilities.write_detail_box_content("<b>Resultat:</b> för närvarande sker kontrollen manuellt, med ovanstående avsnittsinnehåll som underlag")
    print("TKB-granskning klar",datetime.datetime.now().replace(microsecond=0))

    # 2do (senare): kontrollera att det finns V-MIM-tabeller (en gemensam eller en per tjänstekontrakt)
    # 2do (senare): kontrollera att meddelandemodelltabellens attribut mappar mot motsvarande i xsd-schemas

######################################################
##### Privata funktioner (från TKB_inspection.py #####
######################################################
"""def TKB_get_interaction_version(interaction_name):
    version_number = "0"
    #2do: extract version number from interaction paragraph
    searched_paragraph_level = DOCX_document_structure_get_exact_levelvalue(interaction_name)
    #__display_paragraph_text_by_paragraph_level(interaction_name,searched_paragraph_level)
    version_number = DOCX_display_paragraph_text_by_paragraph_level(searched_paragraph_level,interaction_name)

    return version_number"""


"""def TKB_display_paragragh_title(searched_title_name):
    result = True
    result_description = ""
    searched_paragraph_level = DOCX_document_structure_get_exact_levelvalue(searched_title_name)
    if searched_paragraph_level != "":
        #result_description = "OK. TKB (" + searched_paragraph_level + "):  \t" + searched_title_name
        result_description = "TKB (" + searched_paragraph_level + "):  \t" + searched_title_name
        #write_output("OK. TKB (" + searched_paragraph_level + "):  \t" + searched_title_name)
    else:
        result_description = "FEL! " + searched_title_name + " verkar inte vara beskrivet i TKB!"
        #write_output("FEL! " + searched_title_name + " verkar inte vara beskrivet i TKB!")
        result = False
    return result, result_description"""

def __display_paragraph_text_by_paragraph_level(searched_paragraph_level,display_keylevel_text):
    global document_paragraph_index_dict
    previous_key = ""
    for key, value in document_paragraph_index_dict.items():
        if key[0:len(searched_paragraph_level)] == searched_paragraph_level:
            key_level_length = key.find(" ")
            if len(key.strip()) > key_level_length:
                this_key_level = key.strip()[0:key_level_length]
                if this_key_level == previous_key:
                    if display_keylevel_text == True:
                        #write_output("\t" + key.strip()[key_level_length+1:])
                        write_detail_box_content("\t" + key.strip()[key_level_length+1:])
                else:
                    #write_output(key)
                    write_detail_box_content(key)
                previous_key = key.strip()[0:key_level_length]
