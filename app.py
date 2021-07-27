
from docx import Document
from flask import Flask, request    # jsonify
from flask_cors import CORS

import html_dashboard
from html_dashboard import *

import globals
from INFO_inspection_information import *
import io
from repo import *
import requests

##############################
# Startup settings
##############################
# Instantiate the App
app = Flask(__name__)
app.config['JSON_SORT_KEYS'] = False

#pip install -U flask-cors
CORS(app)
app.config['CORS_HEADERS'] = 'Content-Type'


##############################
# App Endpoints
##############################
app = Flask(__name__)
@app.route('/')
def emptyrequest():
    """
    Tar emot anrop för endpoint: "/"

    Sätter ihop och returnerar text som ska underlätta för användaren att förstå hur anropet ska utformas.

    Returnerar: en sträng med html-innehåll
    """
    ##### PREPARE #####
    html = ""
    html = "<br><h1>Webbadressen är inte korrekt!</h1>"
    html += "<br>Någon av de obligatoriska url-parametrarna <i>domain</i> eller <i>tag</i> <b>saknas i anropet!</b>"
    html += "<br><br>Ange dem i adressraden enligt följande format: <i>url...</i><b>/granskningsinfo?domain=</b><i>[domännamn utan riv-prefix]</i><b>&tag=</b><i>[tag]</i>"
    html += "<br><br>Exempelvis: <i><a href='https://i-granskning.herokuapp.com/granskningsinfo?domain=clinicalprocess.healthcond.certificate&tag=4.0.5'>https://i-granskning.herokuapp.com/granskningsinfo?domain=clinicalprocess.healthcond.certificate&tag=4.0.5</a></i>"

    ##### REPLY #####
    return html

@app.route('/granskningsinfo')
#@cross_origin(origin='*',headers=['Content-Type','Authorization'])
def reponse2request():
    """
    Tar emot anrop för endpoint: "/granskningsinfo"

    Verifierar att obligatoriska inparametrar finns, samt anropar då metoder för att granska infospec och TKB

    Returnerar: en sträng med html-innehåll
    """

    ##### PREPARE #####
    globals.GLOBALS_init()
    #detail_box_content = ""
    domain = request.args.get('domain', default="")
    domain = domain.replace("riv.","")
    domain = domain.replace("riv-application.","")
    globals.domain_name = domain
    domain_prefix_param = request.args.get('domainprefix', default="")
    tag = request.args.get('tag', default="")
    globals.tag = tag
    alt_IS_name = request.args.get('is', default="")
    alt_TKB_name = request.args.get('tkb', default="")

    if domain != "" and tag != "":
        if domain_prefix_param.strip() != "":
            riv_domain = "riv-application." + domain
            globals.domain_prefix = "riv-application."
        else:
            riv_domain = "riv."+domain
            globals.domain_prefix = "riv."

        ##### INSPECT #####
        globals.docx_document = globals.IS
        __inspect_IS_document(domain, tag, alt_IS_name)

        globals.docx_document = globals.TKB
        __inspect_TKB_document(domain, tag, alt_TKB_name)

        #html = __get_html_response(riv_domain, IS_page_link, TKB_page_link, IS_document_paragraphs, TKB_document_paragraphs)
        html = html_dashboard.get_page_html()
    else:
        html = "<br><h1>Webbadressen är inte korrekt!</h1>"
        html += "<br>Någon av de obligatoriska url-parametrarna <i>domain</i> eller <i>tag</i> <b>saknas i anropet!</b>"
        html += "<br><br>Ange dem i adressraden enligt följande format: <i>url...</i><b>/granskningsinfo?domain=</b><i>[domännamn utan riv-prefix]</i><b>&tag=</b><i>[tag]</i>"

    ##### REPLY #####
    return html

##############################
# Internal methods
##############################
def __inspect_IS_document(domain, tag, alt_document_name):
    """
    Beräknar url till infospecdokumentet för angiven domain och tag.

    Laddar ner dokumentet till en virtuell fil som läses in i ett docx-Document.

    Anropar därefter metoden "INFO_inspect_document" som genomför granskning av dokumentet.
    """
    global IS_page_link
    global IS_document_paragraphs
    IS_page_link = __get_document_page_link(domain, tag, globals.IS)
    downloaded_IS_page = __get_downloaded_document(IS_page_link)

    IS_document_paragraphs = ""

    IS_head_hash = __get_head_hash(downloaded_IS_page)
    IS_document_link = __get_document_link(domain, tag, globals.IS, IS_head_hash, alt_document_name)
    downloaded_IS_document = __get_downloaded_document(IS_document_link)
    if downloaded_IS_document.status_code == 404:
        ###IS_document_paragraphs = APP_text_document_not_found(globals.IS, domain, tag)
        ###globals.granskningsresultat += "<br><h2>Infospec</h2>" + APP_text_document_not_found(globals.IS, domain, tag)
        #globals.IS_felmeddelande = APP_text_document_not_found(globals.IS, domain, tag)
        globals.IS_exists = False
        docx_IS_document = ""
    else:
        globals.docx_IS_document = __get_docx_document(downloaded_IS_document)
        globals.IS_document_exists = True
        globals.IS_exists = True
        ### dev test ###
        for paragraph in globals.docx_IS_document.paragraphs:
            if paragraph.text.strip() != "":
                IS_document_paragraphs += paragraph.text + "<br>"
        ### dev test ###
        INFO_inspect_document(globals.IS)


def __inspect_TKB_document(domain, tag, alt_document_name):
    """
    Beräknar url till TKB-dokumentet för angiven domain och tag.

    Laddar ner dokumentet till en virtuell fil som läses in i ett docx-Document.

    Anropar därefter metoden "INFO_inspect_document" som genomför granskning av dokumentet.
    """
    global TKB_page_link
    global TKB_document_paragraphs
    TKB_page_link = __get_document_page_link(domain, tag, globals.TKB)
    downloaded_TKB_page = __get_downloaded_document(TKB_page_link)

    TKB_document_paragraphs = ""

    TKB_head_hash = __get_head_hash(downloaded_TKB_page)
    TKB_document_link = __get_document_link(domain, tag, globals.TKB, TKB_head_hash, alt_document_name)
    downloaded_TKB_document = __get_downloaded_document(TKB_document_link)
    if downloaded_TKB_document.status_code == 404:
        ###TKB_document_paragraphs = APP_text_document_not_found(globals.TKB, domain, tag)
        ###globals.granskningsresultat += "<br><br><h2>TKB</h2>" + APP_text_document_not_found(globals.TKB, domain, tag)
        docx_TKB_document = ""
        #globals.TKB_felmeddelande = APP_text_document_not_found(globals.TKB, domain, tag)
        globals.TKB_exists = False
    else:
        globals.docx_TKB_document = __get_docx_document(downloaded_TKB_document)
        globals.TKB_document_exists = True
        globals.TKB_exists = True
        ### dev test ###
        for paragraph in globals.docx_TKB_document.paragraphs:
            if paragraph.text.strip() != "":
                TKB_document_paragraphs += paragraph.text + "<br>"
        ### dev test ###
        INFO_inspect_document(globals.TKB)


def __get_document_page_link(domainname, tag, document):
    """
    Beräknar url till sidan som innehåller länk till angivet dokument för vald domän och tag i Bitbucket-repot.

    Returnerar: länk till dokumentsidan
    """
    url_prefix = "https://bitbucket.org/rivta-domains/"
    url_domain = globals.domain_prefix + domainname + "/"
    url_src = "src/"
    url_tag = tag + "/"
    url_docs = "docs/"
    domain_name = domainname.replace(".","_")
    url_doc = document +"_" + domain_name + ".docx"
    document_page_link = url_prefix+url_domain+url_src+url_tag+url_docs+url_doc

    return document_page_link

"""def __get_domain_docs_link(domainname, tag):
    #
    #Beräknar url till docs-sidan för vald domän och tag i Bitbucket-repot.
    #
    #Returnerar: länk till dokumentsidan
    #
    url_prefix = "https://bitbucket.org/rivta-domains/"
    url_domain = globals.domain_prefix + domainname + "/"
    url_src = "src/"
    url_tag = tag + "/"
    url_docs = "docs/"
    document_page_link = url_prefix+url_domain+url_src+url_tag+url_docs

    return document_page_link"""

def __get_document_link(domainname, tag, document, head_hash, alt_document_name):
    """
    Beräknar url till angivet dokument för vald domän och tag i Bitbucket-repot.

    Returnerar: länk som kan användas vid nerladdning av angivet dokument
    """
    url_prefix = "https://bitbucket.org/rivta-domains/"
    url_domain = globals.domain_prefix + domainname + "/"
    url_raw = "raw/"
    url_docs = "docs/"
    domain_name = domainname.replace(".","_")
    if len(alt_document_name.strip()) > 0:
        url_doc = alt_document_name
    else:
        url_doc = document +"_" + domain_name + ".docx"
    document_link = url_prefix+url_domain+url_raw+head_hash+"/"+url_docs+url_doc

    if document == globals.IS:
        globals.IS_document_name = url_doc
    elif document == globals.TKB:
        globals.TKB_document_name = url_doc

    return document_link

def __get_downloaded_document(document_link):
    """
    Laddar ner dokument från angiven länk.

    Returnerar: nerladdat dokument
    """
    downloaded_doc = requests.get(document_link, stream=True)

    return downloaded_doc

"""def __get_document_in_docx_format(document):
    docx_document = document.content

    return docx_document"""

def __get_head_hash(document_page):
    """
    hämtar head-hash för det dokument som ska laddas ner. Hashen finns i den Bitbucketsida som innehåller länk till dokumentet.

    Returnerar: head-hash
    """
    hash_start = document_page.text.find('{"hash":')
    hash_end = hash_start+17
    head_hash = document_page.text[hash_start+10:hash_end]

    return head_hash

def __get_docx_document(downloaded_document):
    """
    Läser in angivet dokuments innehåll i ett docx-Document.

    Returnerar: docx-Documentet
    """
    with io.BytesIO(downloaded_document.content) as inmemoryfile:
        docx_document = Document(inmemoryfile)

    return docx_document

#def APP_text_document_not_found(doc, domain, tag):

    """
    Sammanställer ett meddelande till användaren då sökt dokument saknas eller då fel dokumentnamn har angivits.

    Returnerar: information i html-format
    """
    """document_name = "Infospec"
    if doc == globals.TKB:
        document_name = globals.TKB

    document_info = "<div><li>"
    document_info += document_name + " saknas eller har annat namn än det förväntade: <i>" + doc.upper() + "_" + domain.replace(".", "_") + ".docx</i>"
    docs_link = REPO_get_domain_docs_link(domain, tag)
    document_info += "<br>Kontrollera dokumentnamn här: <a href='" + docs_link + "'" + " target='_blank'>" + docs_link + "</a>"
    document_info += "<br>Om det finns en " + document_name + " så har den ett annat än det förväntade namnet. "
    document_info += "I så fall kan du ange det namnet som en url-parameter enligt: <i>url...</i><b>&is=dokumentnamn</b>"
    document_info += "<br>Om detta är en applikationsspecifik domän kan du ange det i en url-parameter enligt: <i>url...</i><b>&domainprefix=true</b>"
    document_info += "<div><li>"

    return document_info"""


#def __get_html_response(riv_domain, IS_page_link, TKB_page_link, IS_document_paragraphs, TKB_document_paragraphs):
    """
    Sammanställer ett meddelande till användaren med granskningsresultat

    Returnerar: information i html-format
    """
    """html = '''
        <h1>I-granskningsstöd för: {}</h1>
        <br><h2><a href={}>Infospec-sida</a> &nbsp;&nbsp;&nbsp; <a href={}>TKB-sida</a></h2>
        <br><h2><b>IS-paragrafer:</b></h2> {}
        <br><h2><b>TKB-paragrafer:</b></h2> {}
        '''.format(riv_domain, IS_page_link, TKB_page_link, IS_document_paragraphs, TKB_document_paragraphs)"""

    """html = '''
        <h1>I-granskningsstöd för: {}</h1>
        <br>{}
        <br>{}
        '''.format(riv_domain, __get_summary_in_html_format(), globals.granskningsresultat)

    return html"""

def __get_summary_in_html_format():
    """
    Sammanställer övergripande information om resultatet av granskning av Infospec och TKB.

    Returnerar: en sträng med html-innehåll
    """
    summary = ""
    if globals.IS_exists == False and globals.TKB_exists == False:
        summary = "<h1>Sammanfattning</h1>"
        summary += "<h2>Inga granskningskontroller har utförts på grund av att varken Infospec eller TKB kunnat hittas.</h2><br>"
    else:
        summary += "<h1>Sammanfattning</h1>"
        if globals.IS_exists == True:
            summary += "<h2>Granskningsresultat för Infospec</h2>"
            if globals.IS_antal_brister_revisionshistorik == 0:
                summary += "Revisionshistoriken är <b>korrekt</b>"
            else:
                summary += "<b>Fel versionsnummer</b> angivet i revisionshistoriken"
            summary += "<br><b>" + str(globals.IS_antal_brister_referenslänkar) + "</b> felaktiga länkar i referenstabellen"
            summary += "<br><b>" + str(globals.IS_antal_brister_klassbeskrivning) + "</b> saknade klassbeskrivningar"
            summary += "<br><b>" + str(globals.IS_antal_brister_multiplicitet) + "</b> saknade multipliciteter i klasstabeller"
            summary += "<br><b>" + str(globals.IS_antal_brister_datatyper) + "</b> fel för datatyper"
            summary += "<br><b>" + str(globals.IS_antal_brister_referensinfomodell) + "</b> saknade referenser till RIM i klasstabeller"
            summary += "<br><b>" + str(globals.IS_antal_brister_tomma_tabellceller) + "</b> tomma celler i klasstabeller"
            summary += "<br>"
        #else:
            #summary += __text_document_not_found(globals.IS,domain,globals.tag)
        if globals.TKB_exists == True:
            summary += "<h2>Granskningsresultat för TKB</h2>"
            if globals.TKB_antal_brister_revisionshistorik == 0:
                summary += "Revisionshistoriken är <b>korrekt</b>"
            else:
                summary += "<b>Fel versionsnummer</b> angivet i revisionshistoriken"


    return summary

if __name__ == '__main__':
    from argparse import ArgumentParser

    port = 4001
    usedHost = 'https://callistabackend.herokuapp.com'
    instance_address = "http://" + usedHost + ":" + str(port)

    usedHost = '127.0.0.1'
    app.run(host=usedHost, port=port)

