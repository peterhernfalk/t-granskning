
import globals
import granskning_AB
import granskning_TKB
import re
import requests.exceptions
from requests import head
from requests.exceptions import MissingSchema
import urllib3


#import glob
#import os
#import os.path
#from requests import head

#local_test = False


### Not in use yet ###
def check_if_file_exists(path, search_pattern):
    file_exist = False
    """if os.path.exists(path):
        os.chdir(path)
        for file in glob.glob(search_pattern):
            file_exist = True"""

    if search_pattern == globals.IS:
        file_exist = globals.IS_exists
    elif search_pattern == globals.TKB:
        file_exist = globals.TKB_exists

    return file_exist

"""def extract_urls_from_table(document, table_number):
    table = document.tables[table_number]
    links = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                xml = paragraph.paragraph_format.element.xml
                xml_str = str(xml)
                if "<w:hyperlink" in xml_str or 'w:val="Hyperlink"' in xml_str:
                    wt_list = re.findall('<w:t[\S\s]*?</w:t>', xml_str)
                    if "<w:t xml:" not in str(wt_list[0]):
                        hyperlink = str(wt_list[0]).replace("<w:t>","").replace("</w:t>","")
                        links.append(hyperlink)
    return links
"""
def extract_urls_from_table(document, table_number):
    """
    Utredning inför kodning:
        "Ärendehantering" in xml_str, "<w:hyperlink" in xml_str

        Länk som innehåller url:
            False True
            <w:hyperlink r:id="rId15" w:history="1">
                <w:r w:rsidR="003441E4" w:rsidRPr="00394169">
                    <w:t>http://rivta.se/documents/ARK_0040/</w:t>

        Länk som innehåller display text, men ingen url:
            True True
            <w:hyperlink r:id="rId14" w:history="1">
                <w:r w:rsidR="00D35DA5">
                    <w:t>Ärendehantering</w:t>
            I document.part.rels:
                Länk-id rId14 innehåller url i rels[rel]._target
    """

    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    rels = document.part.rels
    document_relations = {}
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            document_relations[rel] = rels[rel]._target

    table = document.tables[table_number]
    links = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                xml_str = str(paragraph.paragraph_format.element.xml)
                if "<w:hyperlink" in xml_str or 'w:val="Hyperlink"' in xml_str:
                    wt_list = re.findall('<w:t[\S\s]*?</w:t>', xml_str)
                    if "<w:t xml:" not in str(wt_list[0]):
                        hyperlink = str(wt_list[0]).replace("<w:t>","").replace("</w:t>","")
                        if hyperlink.endswith("%20"):
                            hyperlink = hyperlink.removesuffix("%20")
                        if "http" not in hyperlink:
                            ### Workaround for w:t that contains disply text instead of url ###
                            relation_id = re.findall('r:id="[\S\s]*?"', xml_str)
                            if len(relation_id) > 0:
                                relation_id_key = str(relation_id[0]).replace('r:id=','').replace('"','')
                                hyperlink = document_relations[relation_id_key]
                            else:
                                link_in_instrtext = re.findall('HYPERLINK "[\S\s]*?"', xml_str)
                                link_in_instrtext_str = str(link_in_instrtext[0]).replace('HYPERLINK "', '').replace('"', '')
                                hyperlink = link_in_instrtext_str
                        """if globals.docx_document == globals.TKB:
                            if "http" not in hyperlink:
                                ### Workaround for when complete url is found in <w:instrText instead of in <w:t ###
                                link_in_instrtext = re.findall('HYPERLINK "[\S\s]*?"', xml_str)
                                link_in_instrtext_str = str(link_in_instrtext[0]).replace('HYPERLINK "','').replace('"','')
                                hyperlink = link_in_instrtext_str"""
                        links.append(hyperlink)
                elif paragraph.text.lower().find("http") >= 0:
                    paragraph_links = paragraph.text.split("\n")
                    for paragraph_link in paragraph_links:
                        if paragraph_link.lower().find("http") >= 0:
                            if paragraph_link.strip() != "":
                                links.append(paragraph_link)

    return links

"""def write_output(text):
    #print(text)
    globals.granskningsresultat += "<br>" + text
    #if globals.docx_document == globals.AB:
    #    print("write_output:", text)
"""
"""if globals.output_channel == globals.CONSOLE:
        print(text)
    elif globals.output_channel == globals.CONSOLE_AND_TEXTFILE:
        print(text)
        f = open(globals.localGitRepo+"/../Granskningar/"+globals.domain+"-"+globals.tag+"-utdata.txt", "a+")
        f.write("\n"+text)
        f.close()
    elif globals.output_channel == globals.HTML:
        print(text)
    elif globals.output_channel == globals.PDF:
        print(text)
    elif globals.output_channel == globals.WORD:
        print(text)
"""

def write_detail_box_content(text):
    """#formatted_text = ''.join(text).replace("\n","<br>")
    detail_box_content = "<li>" + text + "</li>"
    if globals.docx_document == globals.TKB:
        globals.TKB_detail_box_contents += detail_box_content
    elif globals.docx_document == globals.AB:
        globals.AB_detail_box_contents += detail_box_content"""
    detail_box_content = "<li>" + text + "</li>"
    if globals.docx_document == globals.TKB:
        granskning_TKB.TKB_detail_box_contents += detail_box_content
    elif globals.docx_document == globals.AB:
        granskning_AB.AB_detail_box_contents += detail_box_content

"""def write_detail_box_html(html_text):
    formatted_text = html_text.replace("\n","<br>")
    detail_box_content = "<li>" + formatted_text + "</li>"
    if globals.docx_document == globals.IS:
        globals.IS_detail_box_contents += formatted_text
    elif globals.docx_document == globals.TKB:
        globals.TKB_detail_box_contents += formatted_text
    elif globals.docx_document == globals.AB:
        globals.AB_detail_box_contents += formatted_text"""


def write_output_without_newline(text):
    formatted_text = text.replace("\n","<br>")
    #print(globals.docx_document,"formatted_text:",formatted_text)
    #print(text, "\t", end="")
    globals.granskningsresultat += formatted_text
    #if globals.output_channel == globals.CONSOLE:
    #    print(text,"\t",end="")


def verify_url_exists(searched_url):
    """status_code = 200
    try:
        response = head(searched_url.strip())
        status_code = response.status_code
    except requests.exceptions.ConnectionError:
        status_code = 400  # http status code, meaning bad request
    except MissingSchema:
        status_code = 400   # http status code, meaning bad request
    except requests.exceptions.InvalidSchema:
        status_code = 400   # http status code, meaning bad request
    #else:
    #    status_code = 400  # http status code, meaning bad request
    return status_code"""
    status_code = 200
    urllib3.disable_warnings()
    try:
        response = head(searched_url.strip(), allow_redirects=True, verify=False)    #timeout=2
        status_code = response.status_code
    except requests.exceptions.ConnectionError:
        status_code = 400  # http status code, meaning bad request
    except MissingSchema:
        status_code = 400   # http status code, meaning bad request
    except requests.exceptions.InvalidSchema:
        status_code = 400   # http status code, meaning bad request
    #else:
    #    status_code = 400  # http status code, meaning bad request
    return status_code

"""if local_test == True:
    searched_url_arrray = []

    # Various URL:s
    searched_url_arrray.append("http://www.example.com")
    searched_url_arrray.append("https://www.marvel.com")
    searched_url_arrray.append("https://www.inera.se/om-inera/for-media/kontakt")
    searched_url_arrray.append("https://www.inera.se/om-inera/for-media/kontak")

    # TKB: Intygstjänster
    searched_url_arrray.append("http://rivta.se")
    searched_url_arrray.append("http://rivta.se/domains/clinicalprocess_healthcond_certificate.html")
    searched_url_arrray.append("http://rivta.se/domains/clinicalprocess_healthcond_certificate.html")
    searched_url_arrray.append("https://inera.atlassian.net/wiki/spaces/KINT/pages/3615655/Kodverk+i+nationella+tj+nstekontrakt")
    searched_url_arrray.append("http://rivta.se/documents/ARK_0001/")

    # IS: Intygstjänster
    searched_url_arrray.append("http://rivta.se/")
    searched_url_arrray.append("http://informationsstruktur.socialstyrelsen.se/")
    searched_url_arrray.append("https://inera.atlassian.net/wiki/spaces/KINT/pages/3615655/Kodverk+i+nationella+tj+nstekontrakt")
    searched_url_arrray.append("http://termbank.socialstyrelsen.se/")
    searched_url_arrray.append("http://www.rivta.se/domains/clinicalprocess_healthcond_certificate.html")
    searched_url_arrray.append("https://inera.atlassian.net/wiki/spaces/EIT/overview")
    searched_url_arrray.append("http://www.rikstermbanken.se/mainMenu.html")
    searched_url_arrray.append("https://www.sis.se/produkter/informationsteknik-kontorsutrustning/ittillampningar/halso-och-sjukvardsinformatik/sseniso139402016/ ")
    searched_url_arrray.append("https://www.sis.se/produkter/informationsteknik-kontorsutrustning/ittillampningar/halso-och-sjukvardsinformatik/sseniso210902011/")

    # IS: IS_clinicalprocess_activity_actions.docx
    searched_url_arrray.append("http://www.inera.se/TJANSTER--PROJEKT/Arkitektur-och-regelverk/")
    searched_url_arrray.append("http://www.socialstyrelsen.se/nationellehalsa/nationellinformationsstruktur")
    searched_url_arrray.append("http://browser.ihtsdotools.org")
    searched_url_arrray.append("http://rivta.se/domains/strategicresourcemanagement_persons_person.html")
    searched_url_arrray.append("https://www.inera.se/kundservice/dokument-och-lankar/tjanster/hsa/hsa-kodverk/")
    searched_url_arrray.append("http://www.socialstyrelsen.se/register/verksamhetpersonal/legitimeradpersonal(hosp)")

    # IS: riv.clinicalprocess.activity.request
    searched_url_arrray.append("https://www.inera.se/digitalisering/arkitektur/")
    searched_url_arrray.append("http://www.socialstyrelsen.se/nationellehalsa/nationellinformationsstruktur")
    searched_url_arrray.append("https://inera.atlassian.net/wiki/spaces/KINT/pages/3615655/Kodverk+i+nationella+tj+nstekontrakt")
    searched_url_arrray.append("http://browser.ihtsdotools.org/")

    # IS: riv.clinicalprocess.logistics.cervixscreening
    searched_url_arrray.append("https://www.inera.se/digitalisering/arkitektur/")
    searched_url_arrray.append("https://www.socialstyrelsen.se/utveckla-verksamhet/e-halsa/nationell-informationsstruktur/")
    searched_url_arrray.append("https://inera.atlassian.net/wiki/spaces/KINT/pages/3615655/Kodverk+i+nationella+tj+nstekontrakt ")
    searched_url_arrray.append("http://termbank.socialstyrelsen.se/ ")
    searched_url_arrray.append("https://browser.ihtsdotools.org/?")
    searched_url_arrray.append("https://www.socialstyrelsen.se/utveckla-verksamhet/e-halsa/snomed-ct/snomed-ct-licens/")
    searched_url_arrray.append("https://www.cancercentrum.se/samverkan/vara-uppdrag/prevention-och-tidig-upptackt/gynekologisk-cellprovskontroll/vardprogram/ ")

    for index in range(len(searched_url_arrray)):
        this_url = searched_url_arrray[index]
        result_code = verify_url_exists(this_url)
        if result_code < 404:
            #print("OK, statuskod: (" + str(result_code) + ") " + this_url)
            print("OK: (" + str(result_code) + ") " + this_url)
        else:
            print("Sidan saknas! Statuskod: (" + str(result_code) + ") " + this_url)"""
